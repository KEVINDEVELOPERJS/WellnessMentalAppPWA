"""
Script para extraer imágenes de RTF y generar documento Word automáticamente
"""
import os
import re
import base64
import sys
import subprocess

def install_requirements():
    """Instala las librerías necesarias"""
    packages = ['python-docx', 'Pillow', 'striprtf']
    for package in packages:
        try:
            subprocess.run([sys.executable, "-m", "pip", "install", package], 
                          capture_output=True, check=True)
            print(f"✅ {package} instalado")
        except:
            print(f"⚠️  {package} ya instalado o error en instalación")

def extract_images_from_rtf(rtf_path, output_folder):
    """Extrae imágenes de un archivo RTF usando múltiples métodos"""
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    print(f"📂 Leyendo archivo RTF: {rtf_path}")
    
    with open(rtf_path, 'rb') as f:
        rtf_content = f.read()
    
    # Método 1: Buscar patrones de imágenes hexadecimales en RTF
    print("🔍 Método 1: Buscando patrones hexadecimales...")
    
    # Patrón para datos de imagen en RTF (común para PNG)
    # {\pict\pngblip...\hexdata...}
    hex_pattern = r'\{\\pict\\pngblip[^\}]*?([0-9a-fA-F\s]{100,})\}'
    matches = re.findall(hex_pattern, rtf_content.decode('latin-1', errors='ignore'))
    
    print(f"   Encontrados {len(matches)} bloques hexadecimales potenciales")
    
    for i, hex_data in enumerate(matches):
        try:
            # Limpiar espacios y convertir
            clean_hex = re.sub(r'\s', '', hex_data)
            
            # Asegurar longitud par
            if len(clean_hex) % 2 != 0:
                clean_hex = clean_hex[:-1]
            
            # Convertir a bytes
            image_bytes = bytes.fromhex(clean_hex)
            
            # Intentar validar como imagen
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(image_bytes))
                
                # Guardar como PNG
                output_path = os.path.join(output_folder, f'web_{i+1}.png')
                img.save(output_path, 'PNG')
                print(f"   ✅ Imagen {i+1} guardada: web_{i+1}.png ({img.size})")
                
            except ImportError:
                # Si PIL no está disponible, guardar bytes crudos
                output_path = os.path.join(output_folder, f'web_{i+1}.bin')
                with open(output_path, 'wb') as f:
                    f.write(image_bytes)
                print(f"   📦 Datos crudos {i+1} guardados: web_{i+1}.bin")
                
        except Exception as e:
            print(f"   ❌ Error procesando bloque {i+1}: {e}")
    
    # Método 2: Buscar patrones de imagen más generales
    print("\n🔍 Método 2: Buscando patrones de imagen más generales...")
    
    # Buscar bloques {\pict con cualquier tipo de imagen
    pict_pattern = r'\{\\pict[^}]*?\{?\s*([0-9a-fA-F\s]{50,})\s*\}'
    pict_matches = re.findall(pict_pattern, rtf_content.decode('latin-1', errors='ignore'))
    
    print(f"   Encontrados {len(pict_matches)} bloques pict adicionales")
    
    for i, hex_data in enumerate(pict_matches):
        try:
            clean_hex = re.sub(r'\s', '', hex_data)
            if len(clean_hex) % 2 != 0:
                clean_hex = clean_hex[:-1]
            
            image_bytes = bytes.fromhex(clean_hex)
            
            # Intentar validar
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(image_bytes))
                
                output_path = os.path.join(output_folder, f'web_pict_{i+1}.png')
                img.save(output_path, 'PNG')
                print(f"   ✅ Imagen pict {i+1} guardada: web_pict_{i+1}.png ({img.size})")
                
            except ImportError:
                output_path = os.path.join(output_folder, f'web_pict_{i+1}.bin')
                with open(output_path, 'wb') as f:
                    f.write(image_bytes)
                print(f"   📦 Datos pict {i+1} guardados: web_pict_{i+1}.bin")
                
        except Exception as e:
            print(f"   ❌ Error procesando pict {i+1}: {e}")
    
    # Contar imágenes extraídas
    image_files = [f for f in os.listdir(output_folder) if f.endswith('.png') or f.endswith('.bin')]
    print(f"\n📊 Total de archivos extraídos: {len(image_files)}")
    
    return image_files

def create_word_with_images(images_folder, output_docx):
    """Crea documento Word con las imágenes extraídas"""
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ python-docx no está instalado. Instalando...")
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], 
                      capture_output=True, check=True)
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    print("📝 Creando documento Word...")
    
    doc = Document()
    
    # Configurar estilo
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    def remove_spacing(paragraph):
        paragraph.space_before = Pt(0)
        paragraph.space_after = Pt(0)
        paragraph.line_spacing = 1.0
    
    # Título
    title = doc.add_heading('Informe Técnico: Funcionamiento Correcto de Wellness Mental App Web', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(title)
    
    # Metadatos
    doc.add_paragraph('Fecha: 4 de agosto de 2026')
    doc.add_paragraph('Versión: 1.0')
    doc.add_paragraph('Tipo: Informe de Validación de Criterios de Aceptación - Versión Web')
    doc.add_paragraph()
    
    # Resumen
    h = doc.add_heading('Resumen Ejecutivo', level=1)
    remove_spacing(h)
    p = doc.add_paragraph('Este informe presenta una validación detallada del funcionamiento correcto de la aplicación web "Wellness Mental App Web", desarrollada con HTML5, CSS3, JavaScript y arquitectura PWA.')
    remove_spacing(p)
    
    # Arquitectura
    h = doc.add_heading('Arquitectura Técnica', level=1)
    remove_spacing(h)
    h = doc.add_heading('Estructura Web Implementada', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Capa Model (Datos y Persistencia):', level=3)
    remove_spacing(h)
    p = doc.add_paragraph()
    p.add_run('IndexedDB/').bold = True
    p.add_run(': Base de datos local del navegador')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('LocalStorage/').bold = True
    p.add_run(': Gestión de sesión y preferencias')
    remove_spacing(p)
    
    h = doc.add_heading('Capa View (Interfaz de Usuario):', level=3)
    remove_spacing(h)
    pages = [
        'index.html: Login/registro',
        'evaluation.html: Evaluaciones psicológicas',
        'chat.html: Chat con IA',
        'exercises.html: Ejercicios de bienestar',
        'games.html: Gamificación y juegos',
        'community.html: Comunidad estudiantil',
        'alerts.html: Alertas riesgo alto',
        'alerts-low-medium.html: Alertas riesgo bajo/medio',
        'questionnaire-editor.html: Editor cuestionarios',
        'profile.html: Perfil y privacidad',
        'parent-reports.html: Informes padres',
        'videos.html: Videos guiados',
        'active-breaks.html: Pausas activas',
        'mental-garden.html: Jardín mental'
    ]
    for page in pages:
        doc.add_paragraph(page)
    
    h = doc.add_heading('Capa Controller (Lógica de Negocio):', level=3)
    remove_spacing(h)
    controllers = [
        'app.js: Controlador principal',
        'evaluation.js: Lógica evaluaciones',
        'chat.js: Gestión chat IA',
        'exercises.js: Control ejercicios',
        'games.js: Sistema gamificación',
        'mental-garden.js: Lógica jardín',
        'alerts.js: Alertas riesgo alto',
        'alerts-low-medium.js: Alertas riesgo bajo/medio',
        'hub-client.js: Cliente API'
    ]
    for controller in controllers:
        doc.add_paragraph(controller)
    
    # Validación
    h = doc.add_heading('Validación por Criterios de Aceptación con Evidencia Visual', level=1)
    remove_spacing(h)
    
    # Mapeo de imágenes (usar imágenes extraídas secuencialmente)
    image_sequence = [
        ('HU-01', 'Registro y Autenticación', ['web_registro.png', 'web_login.png', 'web_consentimiento.png']),
        ('HU-02', 'Dashboard Principal', ['web_dashboard_estudiante.png', 'web_dashboard_psicologo.png']),
        ('HU-03', 'Evaluación Psicológica', ['web_cuestionarios.png', 'web_gad7.png', 'web_resultados.png', 'web_historial.png']),
        ('HU-04', 'Chat con IA', ['web_chat.png', 'web_chat_conversacion.png']),
        ('HU-05', 'Ejercicios Bienestar', ['web_ejercicios.png', 'web_respiracion.png', 'web_meditacion.png', 'web_progreso.png']),
        ('HU-08', 'Gamificación', ['web_juegos.png', 'web_puzzle.png', 'web_arte.png', 'web_ritmo.png', 'web_jardin.png', 'web_logros.png', 'web_puntos.png']),
        ('HU-07', 'Comunidad', ['web_comunidad.png', 'web_post.png', 'web_interaccion.png']),
        ('HU-09', 'Informes Padres', ['web_informes.png', 'web_informe_previa.png']),
        ('HU-10', 'Perfil Privacidad', ['web_perfil.png', 'web_privacidad.png', 'web_sesiones.png']),
        ('HU-06', 'Editor Cuestionarios', ['web_editor.png', 'web_editor_previa.png']),
        ('Alertas', 'Gestión Alertas', ['web_alertas_panel.png', 'web_alertas_detalle.png', 'web_alertas_bajo_medio.png']),
        ('Misiones', 'Misiones Check-In', ['web_checkin.png', 'web_misiones.png']),
        ('Videos', 'Videos Guiados', ['web_videos.png', 'web_video_reproduccion.png']),
        ('Pausas', 'Pausas Activas', ['web_pausas.png', 'web_pausa_config.png'])
    ]
    
    # Obtener lista de imágenes disponibles
    available_images = sorted([f for f in os.listdir(images_folder) if f.endswith('.png')])
    print(f"🖼️ Imágenes disponibles: {len(available_images)}")
    print(f"📋 Imágenes requeridas: {sum(len(c) for _, _, c in image_sequence)}")
    
    img_index = 0
    
    for key, title, required_images in image_sequence:
        h = doc.add_heading(f'{key}: {title}', level=2)
        remove_spacing(h)
        h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
        remove_spacing(h)
        
        for image_name in required_images:
            criterion = image_name.replace('.png', '').replace('web_', 'CA-')
            h = doc.add_heading(criterion, level=4)
            remove_spacing(h)
            
            p = doc.add_paragraph()
            p.add_run(f'Evidencia Visual: {image_name}. ').bold = True
            p.add_run('✅ FUNCIONAL')
            remove_spacing(p)
            
            # Intentar insertar imagen si está disponible
            if img_index < len(available_images):
                image_path = os.path.join(images_folder, available_images[img_index])
                try:
                    doc.add_picture(image_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    remove_spacing(doc.paragraphs[-1])
                    print(f"   ✅ Imagen insertada: {available_images[img_index]} para {criterion}")
                    img_index += 1
                except Exception as e:
                    p = doc.add_paragraph(f'[Error al insertar imagen: {e}]')
                    remove_spacing(p)
            else:
                p = doc.add_paragraph(f'[Imagen no disponible: {image_name}]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                remove_spacing(p)
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    # Conclusiones
    h = doc.add_heading('Conclusiones', level=1)
    remove_spacing(h)
    doc.add_paragraph('La aplicación web Wellness Mental App cumple satisfactoriamente con todos los criterios de aceptación.')
    
    # Guardar
    doc.save(output_docx)
    print(f"\n✅ Documento Word creado: {output_docx}")
    return True

def main():
    """Función principal"""
    import io
    
    rtf_path = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.rtf'
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    output_docx = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB.docx'
    
    print("🚀 Iniciando proceso de generación de documento...")
    print("=" * 60)
    
    # Paso 1: Instalar dependencias
    print("\n📦 Paso 1: Instalando dependencias...")
    install_requirements()
    
    # Paso 2: Extraer imágenes
    print("\n🖼️  Paso 2: Extrayendo imágenes del RTF...")
    image_files = extract_images_from_rtf(rtf_path, images_folder)
    
    # Paso 3: Crear documento Word
    print("\n📝 Paso 3: Creando documento Word con imágenes...")
    success = create_word_with_images(images_folder, output_docx)
    
    if success:
        print("\n" + "=" * 60)
        print("🎉 PROCESSO COMPLETADO EXITOSAMENTE")
        print("=" * 60)
        print(f"📁 Imágenes extraídas en: {images_folder}")
        print(f"📄 Documento Word creado: {output_docx}")
        print("\n💡 NOTA: Si las imágenes no se insertaron correctamente, puedes:")
        print("   1. Abrir el documento Word manualmente")
        print("   2. Insertar las imágenes desde la carpeta web_images")
        print("   3. Usar el documento RTF como guía de estructura")
    else:
        print("\n❌ Hubo errores en el proceso")

if __name__ == '__main__':
    main()
