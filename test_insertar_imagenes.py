from docx import Document
from docx.shared import Inches
import os

def test_image_insertion():
    """Prueba simple de inserción de imágenes"""
    doc = Document()
    
    # Carpeta de imágenes
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    # Obtener primeras 5 imágenes
    image_files = []
    if os.path.exists(images_folder):
        import glob
        image_files = sorted(glob.glob(os.path.join(images_folder, 'image*.png')))[:5]
    
    print(f"Imágenes encontradas: {len(image_files)}")
    for img in image_files:
        print(f"  - {os.path.basename(img)}")
    
    # Insertar imágenes en el documento
    for i, img_path in enumerate(image_files):
        try:
            doc.add_heading(f'Imagen {i+1}', level=2)
            doc.add_picture(img_path, width=Inches(4))
            doc.add_paragraph(f"Ruta: {img_path}")
            print(f"Imagen {i+1} insertada correctamente")
        except Exception as e:
            doc.add_paragraph(f"Error insertando imagen {i+1}: {e}")
            print(f"Error insertando imagen {i+1}: {e}")
    
    # Guardar documento de prueba
    output_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\test_imagenes.docx'
    doc.save(output_path)
    print(f"Documento de prueba guardado en: {output_path}")

if __name__ == "__main__":
    test_image_insertion()
