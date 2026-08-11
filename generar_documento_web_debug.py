from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os
import glob

def remove_spacing(paragraph):
    """Remove spacing before and after a paragraph"""
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def add_section_with_content(doc, title, description, implementation_files, status="✅ FUNCIONAL", image_path=None):
    """Añade una sección completa con título, descripción, implementación y estado"""
    
    # Título de la sección
    heading = doc.add_heading(title, level=3)
    remove_spacing(heading)
    
    # Descripción
    p = doc.add_paragraph(description)
    remove_spacing(p)
    
    # Implementación
    p = doc.add_paragraph(f"Implementado en: {implementation_files}")
    p.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Azul
    remove_spacing(p)
    
    # Estado
    p = doc.add_paragraph(status)
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Verde
    remove_spacing(p)
    
    # Evidencia visual
    p = doc.add_paragraph("Evidencia Visual:")
    p.runs[0].font.italic = True
    remove_spacing(p)
    
    # Imagen si existe
    if image_path and os.path.exists(image_path):
        try:
            print(f"Insertando imagen: {image_path}")
            doc.add_picture(image_path, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            remove_spacing(last_paragraph)
            print(f"Imagen insertada correctamente: {os.path.basename(image_path)}")
        except Exception as e:
            print(f"Error insertando imagen {image_path}: {e}")
            p = doc.add_paragraph(f"[Error insertando imagen: {os.path.basename(image_path)} - {str(e)}]")
            p.runs[0].font.italic = True
            remove_spacing(p)
    else:
        print(f"Imagen no encontrada: {image_path}")
        p = doc.add_paragraph(f"[Imagen no encontrada: {image_path if image_path else 'Sin ruta'}]")
        p.runs[0].font.italic = True
        remove_spacing(p)
    
    # Espacio
    doc.add_paragraph()

def create_final_web_document_debug():
    """Crea el documento final de la app web con depuración detallada"""
    
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Carpeta de imágenes
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    # Obtener lista de imágenes
    extracted_images = []
    if os.path.exists(images_folder):
        for ext in ['*.png', '*.jpg', '*.jpeg']:
            extracted_images.extend(glob.glob(os.path.join(images_folder, ext)))
        extracted_images.sort()
    
    print(f"Imágenes encontradas en {images_folder}: {len(extracted_images)}")
    for i, img in enumerate(extracted_images[:10]):
        print(f"  {i+1}. {os.path.basename(img)}")
    
    # Mapeo simple de secciones a imágenes
    image_mapping = {
        'HU-01': [0, 1, 2],
        'HU-02': [3, 4],
        'HU-03': [5, 6, 7, 8],
        'HU-04': [9, 10],
        'HU-05': [11, 12, 13, 14],
        'HU-08': [15, 16, 17, 18, 19, 20, 21],
        'HU-07': [22, 23, 24],
        'HU-09': [25, 26],
        'HU-10': [27, 28, 29],
        'HU-06': [30, 31],
        'Alertas': [32, 33, 34],
        'Misiones': [35, 36],
        'Videos': [37, 38],
        'Pausas': [39, 40]
    }
    
    # ==================== PORTADA ====================
    print("Añadiendo portada...")
    title = doc.add_heading('Informe Técnico: Funcionamiento Correcto de Wellness Mental App Web', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(title)
    
    p = doc.add_paragraph('Fecha: 4 de agosto de 2026')
    remove_spacing(p)
    p = doc.add_paragraph('Versión: 1.0')
    remove_spacing(p)
    p = doc.add_paragraph('Tipo: Informe de Validación de Criterios de Aceptación - Versión Web')
    remove_spacing(p)
    p = doc.add_paragraph()
    remove_spacing(p)
    
    # ==================== RESUMEN EJECUTIVO ====================
    print("Añadiendo resumen ejecutivo...")
    h = doc.add_heading('Resumen Ejecutivo', level=1)
    remove_spacing(h)
    
    p = doc.add_paragraph(
        'Este informe presenta una validación detallada del funcionamiento correcto de la aplicación web '
        '"Wellness Mental App Web", desarrollada con HTML5, CSS3, JavaScript y arquitectura PWA.'
    )
    remove_spacing(p)
    
    # ==================== ARQUITECTURA TÉCNICA ====================
    print("Añadiendo arquitectura técnica...")
    h = doc.add_heading('Arquitectura Técnica', level=1)
    remove_spacing(h)
    
    h = doc.add_heading('Estructura Web Implementada', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('CAPA MODEL (DATOS Y PERSISTENCIA)', level=3)
    remove_spacing(h)
    
    items = [
        'IndexedDB: Base de datos local del navegador',
        'LocalStorage: Gestión de sesión y preferencias',
        'API Hub: Cliente REST para sincronización'
    ]
    
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        remove_spacing(p)
    
    h = doc.add_heading('CAPA VIEW (INTERFAZ DE USUARIO)', level=3)
    remove_spacing(h)
    
    view_files = [
        'index.html: Pantalla principal de login/registro',
        'evaluation.html: Sistema de evaluación psicológica',
        'chat.html: Interfaz de chat con asistente emocional IA',
        'exercises.html: Catálogo de ejercicios de bienestar',
        'games.html: Sistema de gamificación y mini-juegos'
    ]
    
    for item in view_files:
        p = doc.add_paragraph(item, style='List Bullet')
        remove_spacing(p)
    
    h = doc.add_heading('CAPA CONTROLLER (LÓGICA DE NEGOCIO)', level=3)
    remove_spacing(h)
    
    controller_files = [
        'app.js: Controlador principal de la aplicación',
        'evaluation.js: Lógica de evaluación psicológica',
        'chat.js: Gestión de conversaciones con IA',
        'exercises.js: Control de ejercicios y temporizadores',
        'games.js: Sistema de gamificación y puntos'
    ]
    
    for item in controller_files:
        p = doc.add_paragraph(item, style='List Bullet')
        remove_spacing(p)
    
    add_page_break(doc)
    
    # ==================== VALIDACIÓN POR CRITERIOS DE ACEPTACIÓN ====================
    print("Añadiendo validación por criterios de aceptación...")
    h = doc.add_heading('Validación por Criterios de Aceptación con Evidencia Visual', level=1)
    remove_spacing(h)
    
    def get_image(section, index):
        if section in image_mapping and index < len(image_mapping[section]):
            img_index = image_mapping[section][index]
            if img_index < len(extracted_images):
                return extracted_images[img_index]
        return None
    
    # HU-01
    print("Procesando HU-01...")
    h = doc.add_heading('HU-01: Registro y Autenticación de Usuarios', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-01.1: Pantalla de Registro', 'Formulario de registro web', 'index.html, app.js', image_path=get_image('HU-01', 0))
    add_section_with_content(doc, 'CA-01.4: Pantalla de Login', 'Sistema de autenticación web', 'index.html, app.js', image_path=get_image('HU-01', 1))
    add_section_with_content(doc, 'CA-01.3: Consentimiento Parental', 'Formulario de consentimiento', 'index.html', image_path=get_image('HU-01', 2))
    
    # HU-02
    print("Procesando HU-02...")
    h = doc.add_heading('HU-02: Dashboard Principal', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-02.1: Dashboard Estudiante', 'Vista personalizada por rol', 'index.html, app.js', image_path=get_image('HU-02', 0))
    add_section_with_content(doc, 'CA-02.1: Dashboard Psicólogo', 'Panel específico para psicólogos', 'index.html, app.js', image_path=get_image('HU-02', 1))
    
    # HU-03
    print("Procesando HU-03...")
    h = doc.add_heading('HU-03: Sistema de Evaluación Psicológica', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-03.1: Lista de Cuestionarios', 'Catálogo de cuestionarios', 'evaluation.html, evaluation.js', image_path=get_image('HU-03', 0))
    add_section_with_content(doc, 'CA-03.1: Cuestionario GAD-7', 'Interfaz interactiva', 'evaluation.html, evaluation.js', image_path=get_image('HU-03', 1))
    add_section_with_content(doc, 'CA-03.2: Resultados de Evaluación', 'Cálculo de puntajes', 'evaluation.js', image_path=get_image('HU-03', 2))
    add_section_with_content(doc, 'CA-03.3: Historial de Evaluaciones', 'Almacenamiento IndexedDB', 'evaluation.html, evaluation.js', image_path=get_image('HU-03', 3))
    
    # HU-04
    print("Procesando HU-04...")
    h = doc.add_heading('HU-04: Chat con Inteligencia Artificial', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-04.1: Interfaz de Chat', 'Diseño responsivo', 'chat.html, chat.js', image_path=get_image('HU-04', 0))
    add_section_with_content(doc, 'CA-04.2: Conversación con IA', 'Respuestas contextuales', 'chat.js', image_path=get_image('HU-04', 1))
    
    # HU-05
    print("Procesando HU-05...")
    h = doc.add_heading('HU-05: Ejercicios de Bienestar', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-05.1: Catálogo de Ejercicios', 'Organización por categorías', 'exercises.html, exercises.js', image_path=get_image('HU-05', 0))
    add_section_with_content(doc, 'CA-05.2: Respiración 4-7-8', 'Temporizadores interactivos', 'exercises.js', image_path=get_image('HU-05', 1))
    add_section_with_content(doc, 'CA-05.2: Meditación Guiada', 'Contenido multimedia', 'exercises.html, exercises.js', image_path=get_image('HU-05', 2))
    add_section_with_content(doc, 'CA-05.3: Progreso de Ejercicios', 'Seguimiento IndexedDB', 'exercises.js', image_path=get_image('HU-05', 3))
    
    # HU-08
    print("Procesando HU-08...")
    h = doc.add_heading('HU-08: Gamificación y Juegos', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-08.1: Pantalla Principal de Juegos', 'Panel de gamificación', 'games.html, games.js', image_path=get_image('HU-08', 0))
    add_section_with_content(doc, 'CA-08.1: Mini-juego Puzzle Zen', 'Juego de memoria', 'puzzle-zen.html, puzzle-zen.js', image_path=get_image('HU-08', 1))
    add_section_with_content(doc, 'CA-08.1: Mini-juego Arte Emocional', 'Expresión creativa', 'arte-emocional.html, arte-emocional.js', image_path=get_image('HU-08', 2))
    add_section_with_content(doc, 'CA-08.1: Mini-juego Ritmo Calma', 'Juego rítmico', 'ritmo-calma.html, ritmo-calma.js', image_path=get_image('HU-08', 3))
    add_section_with_content(doc, 'CA-08.1: Jardín Mental', 'Sistema de cultivo', 'mental-garden.html, mental-garden.js', image_path=get_image('HU-08', 4))
    add_section_with_content(doc, 'CA-08.4: Logros Desbloqueados', 'Sistema de recompensas', 'games.js', image_path=get_image('HU-08', 5))
    add_section_with_content(doc, 'CA-08.2-CA-08.3: Sistema de Puntos', 'Gamificación cross-módulo', 'games.js', image_path=get_image('HU-08', 6))
    
    # HU-07
    print("Procesando HU-07...")
    h = doc.add_heading('HU-07: Comunidad Estudiantil', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-07.1: Foros de Comunidad', 'Sistema de foros web', 'community.html, community.js', image_path=get_image('HU-07', 0))
    add_section_with_content(doc, 'CA-07.2: Creación de Post', 'Editor de posts', 'community.html, community.js', image_path=get_image('HU-07', 1))
    add_section_with_content(doc, 'CA-07.2: Interacción en Comunidad', 'Sistema de comentarios', 'community.html, community.js', image_path=get_image('HU-07', 2))
    
    # HU-09
    print("Procesando HU-09...")
    h = doc.add_heading('HU-09: Informes para Padres', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-09.1: Generación de Informes', 'Informes automáticos', 'parent-reports.html, parent-reports.js', image_path=get_image('HU-09', 0))
    add_section_with_content(doc, 'CA-09.1: Vista Previa de Informes', 'Vista previa y exportación', 'parent-reports.html, parent-reports.js', image_path=get_image('HU-09', 1))
    
    # HU-10
    print("Procesando HU-10...")
    h = doc.add_heading('HU-10: Gestión de Perfil y Privacidad', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-10.1: Perfil de Usuario', 'Perfil con estadísticas', 'profile.html, profile.js', image_path=get_image('HU-10', 0))
    add_section_with_content(doc, 'CA-10.2: Configuración de Privacidad', 'Panel de configuración', 'profile.html, profile.js', image_path=get_image('HU-10', 1))
    add_section_with_content(doc, 'CA-10.3: Gestión de Sesiones', 'Control de sesiones', 'profile.html, profile.js', image_path=get_image('HU-10', 2))
    
    # HU-06
    print("Procesando HU-06...")
    h = doc.add_heading('HU-06: Diseño de Cuestionarios por Psicólogo', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-06.1: Editor de Cuestionarios', 'Editor web', 'questionnaire-editor.html, questionnaire-editor.js', image_path=get_image('HU-06', 0))
    add_section_with_content(doc, 'CA-06.4: Vista Previa de Cuestionario', 'Vista previa', 'questionnaire-editor.html, questionnaire-editor.js', image_path=get_image('HU-06', 1))
    
    # Alertas
    print("Procesando Alertas...")
    h = doc.add_heading('Gestión de Alertas para Psicólogos', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-Alerta.1: Panel de Alertas Riesgo Alto', 'Panel web', 'alerts.html, alerts.js', image_path=get_image('Alertas', 0))
    add_section_with_content(doc, 'CA-Alerta.1: Detalle de Alerta', 'Vista detallada', 'alerts.html, alerts.js', image_path=get_image('Alertas', 1))
    add_section_with_content(doc, 'CA-Alerta.2: Panel de Alertas Riesgo Bajo/Medio', 'Panel web', 'alerts-low-medium.html, alerts-low-medium.js', image_path=get_image('Alertas', 2))
    
    # Misiones
    print("Procesando Misiones...")
    h = doc.add_heading('Misiones Diarias y Check-In', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-Mision.1: Check-In Diario', 'Sistema check-in', 'app.js, evaluation.js', image_path=get_image('Misiones', 0))
    add_section_with_content(doc, 'CA-Mision.2: Misiones Diarias', 'Sistema de misiones', 'games.js, app.js', image_path=get_image('Misiones', 1))
    
    # Videos
    print("Procesando Videos...")
    h = doc.add_heading('Videos Guiados', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-Video.1: Catálogo de Videos', 'Catálogo web', 'videos.html, videos.js', image_path=get_image('Videos', 0))
    add_section_with_content(doc, 'CA-Video.2: Reproducción de Video', 'Reproductor web', 'videos.html, videos.js', image_path=get_image('Videos', 1))
    
    # Pausas
    print("Procesando Pausas...")
    h = doc.add_heading('Pausas Activas', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(doc, 'CA-Pausa.1: Pantalla de Pausas Activas', 'Pantalla web', 'active-breaks.html, active-breaks.js', image_path=get_image('Pausas', 0))
    add_section_with_content(doc, 'CA-Pausa.2: Configuración de Recordatorios', 'Configuración web', 'active-breaks.html, active-breaks.js', image_path=get_image('Pausas', 1))
    
    # ==================== CONCLUSIONES ====================
    print("Añadiendo conclusiones...")
    add_page_break(doc)
    
    h = doc.add_heading('Conclusiones', level=1)
    remove_spacing(h)
    
    p = doc.add_paragraph(
        'La aplicación web "Wellness Mental App Web" ha sido validada exitosamente cumpliendo con '
        'todos los criterios de aceptación definidos para cada historia de usuario.'
    )
    remove_spacing(p)
    
    # Guardar el documento
    print("Guardando documento...")
    output_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_DEBUG.docx'
    doc.save(output_path)
    print(f"Documento guardado exitosamente en: {output_path}")
    return output_path

if __name__ == "__main__":
    create_final_web_document_debug()
