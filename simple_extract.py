"""
Script simplificado para extraer imágenes y generar documento
"""
import os
import re
import sys

def main():
    print("🚀 Iniciando proceso simplificado...")
    
    # Rutas
    rtf_path = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.rtf'
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    # Crear carpeta
    if not os.path.exists(images_folder):
        os.makedirs(images_folder)
        print(f"📁 Carpeta creada: {images_folder}")
    
    # Leer RTF
    print(f"📂 Leyendo archivo RTF...")
    try:
        with open(rtf_path, 'rb') as f:
            rtf_content = f.read()
        print("✅ Archivo RTF leído correctamente")
    except Exception as e:
        print(f"❌ Error leyendo RTF: {e}")
        return
    
    # Decodificar para buscar patrones
    try:
        rtf_text = rtf_content.decode('latin-1', errors='ignore')
    except:
        rtf_text = rtf_content.decode('utf-8', errors='ignore')
    
    # Buscar patrones de imagen
    print("🔍 Buscando patrones de imagen...")
    
    # Patrón para datos hexadecimales largos (posibles imágenes)
    hex_patterns = re.findall(r'([0-9a-fA-F]{100,})', rtf_text)
    
    print(f"📊 Encontrados {len(hex_patterns)} patrones hexadecimales largos")
    
    # Procesar cada patrón
    valid_images = 0
    for i, hex_data in enumerate(hex_patterns[:50]):  # Limitar a primeros 50
        try:
            # Limpiar y convertir
            clean_hex = re.sub(r'[^0-9a-fA-F]', '', hex_data)
            
            if len(clean_hex) < 100:  # Ignorar muy cortos
                continue
                
            if len(clean_hex) % 2 != 0:
                clean_hex = clean_hex[:-1]
            
            # Convertir a bytes
            image_bytes = bytes.fromhex(clean_hex)
            
            # Verificar si parece una imagen PNG (tiene cabecera PNG)
            if len(image_bytes) > 8 and image_bytes[:8] == b'\x89PNG\r\n\x1a\n':
                # Es PNG, guardar
                output_path = os.path.join(images_folder, f'web_{valid_images+1}.png')
                with open(output_path, 'wb') as f:
                    f.write(image_bytes)
                print(f"✅ Imagen PNG {valid_images+1} guardada: web_{valid_images+1}.png")
                valid_images += 1
            elif len(image_bytes) > 4 and image_bytes[:2] == b'BM':
                # Es BMP, guardar como PNG
                try:
                    from PIL import Image
                    import io
                    img = Image.open(io.BytesIO(image_bytes))
                    output_path = os.path.join(images_folder, f'web_{valid_images+1}.png')
                    img.save(output_path, 'PNG')
                    print(f"✅ Imagen BMP convertida {valid_images+1}: web_{valid_images+1}.png")
                    valid_images += 1
                except:
                    # Guardar como binario
                    output_path = os.path.join(images_folder, f'web_{valid_images+1}.bin')
                    with open(output_path, 'wb') as f:
                        f.write(image_bytes)
                    print(f"📦 Datos binarios {valid_images+1}: web_{valid_images+1}.bin")
                    valid_images += 1
            else:
                # Guardar como binario para análisis
                output_path = os.path.join(images_folder, f'web_raw_{valid_images+1}.bin')
                with open(output_path, 'wb') as f:
                    f.write(image_bytes)
                print(f"📦 Datos crudos {valid_images+1}: web_raw_{valid_images+1}.bin")
                valid_images += 1
                
        except Exception as e:
            print(f"❌ Error procesando patrón {i}: {e}")
    
    print(f"\n📊 Total de archivos extraídos: {valid_images}")
    
    # Ahora crear documento Word básico
    print("\n📝 Creando documento Word básico...")
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Título
        title = doc.add_heading('Informe Técnico: Wellness Mental App Web', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadatos
        doc.add_paragraph('Fecha: 4 de agosto de 2026')
        doc.add_paragraph('Versión: 1.0')
        doc.add_paragraph('Tipo: Informe de Validación - Versión Web')
        doc.add_paragraph()
        
        # Resumen
        doc.add_heading('Resumen Ejecutivo', level=1)
        doc.add_paragraph('Informe de validación de la aplicación web Wellness Mental App.')
        
        # Agregar imágenes extraídas
        doc.add_heading('Imágenes Extraídas del RTF', level=1)
        
        image_files = sorted([f for f in os.listdir(images_folder) if f.endswith('.png')])
        print(f"🖼️ Imágenes PNG encontradas: {len(image_files)}")
        
        for img_file in image_files:
            doc.add_paragraph(f'Imagen: {img_file}')
            img_path = os.path.join(images_folder, img_file)
            try:
                doc.add_picture(img_path, width=Inches(5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f'[Error al insertar {img_file}: {e}]')
        
        # Guardar
        output_docx = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_IMAGENES_EXTRAIDAS.docx'
        doc.save(output_docx)
        print(f"✅ Documento Word creado: {output_docx}")
        
    except Exception as e:
        print(f"❌ Error creando documento Word: {e}")
        print("💡 Puedes usar las imágenes extraídas manualmente en el documento RTF")

if __name__ == '__main__':
    main()
