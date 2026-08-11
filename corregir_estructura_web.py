from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import os
import shutil

def remove_spacing(paragraph):
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def correct_web_structure():
    """Corrige la sección de arquitectura técnica para mostrar estructura web correcta"""
    
    input_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_FINAL_COMPLETO.docx'
    output_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_CORREGIDO.docx'
    
    print("=== CORRIGIENDO ESTRUCTURA WEB ===")
    print("1. Copiando documento...")
    shutil.copy2(input_doc, output_doc)
    
    doc = Document(output_doc)
    
    # Buscar y reemplazar la sección de arquitectura técnica
    architecture_found = False
    in_architecture_section = False
    paragraphs_to_remove = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        
        # Detectar inicio de sección de arquitectura
        if 'Arquitectura Técnica' in text or 'ESTRUCTURA' in text:
            architecture_found = True
            in_architecture_section = True
            paragraphs_to_remove.append(i)
            continue
        
        # Si estamos en la sección de arquitectura, marcar para eliminar
        if in_architecture_section:
            if 'Validación' in text or 'Criterios' in text:
                in_architecture_section = False
            else:
                paragraphs_to_remove.append(i)
    
    # Eliminar párrafos de la arquitectura antigua
    for i in reversed(paragraphs_to_remove):
        p = doc.paragraphs[i]
        p_element = p._element
        p_element.getparent().remove(p_element)
    
    # Insertar nueva arquitectura web correcta
    # Buscar posición después del resumen ejecutivo
    insert_position = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if 'Resumen Ejecutivo' in paragraph.text:
            # Insertar después del resumen
            insert_position = i + 1
            break
    
    # Nueva arquitectura web correcta
    new_architecture = [
        ('Arquitectura Técnica', 1),
        ('Estructura Web Implementada', 2),
        ('CAPA MODEL (DATOS Y PERSISTENCIA)', 3),
        ('', 0),  # espacio
        ('IndexedDB: Base de datos local del navegador para almacenamiento persistente (usuarios, respuestas, resultados, chat, ejercicios)', 0),
        ('LocalStorage: Gestión de sesión y preferencias de usuario', 0),
        ('API Hub: Cliente REST para sincronización con backend centralizado', 0),
        ('SessionStorage: Gestión de datos temporales de sesión', 0),
        ('', 0),
        ('CAPA VIEW (INTERFAZ DE USUARIO)', 3),
        ('', 0),
        ('index.html: Pantalla principal de login/registro y dashboard', 0),
        ('evaluation.html: Sistema de evaluación psicológica (GAD-7, PHQ-9, PSS-10)', 0),
        ('chat.html: Interfaz de chat con asistente emocional IA', 0),
        ('exercises.html: Catálogo de ejercicios de bienestar y técnicas de relajación', 0),
        ('games.html: Sistema de gamificación y mini-juegos', 0),
        ('community.html: Foros de comunidad estudiantil', 0),
        ('alerts.html: Panel de alertas para psicólogos (riesgo alto)', 0),
        ('alerts-low-medium.html: Panel de alertas para psicólogos (riesgo bajo/medio)', 0),
        ('questionnaire-editor.html: Editor de cuestionarios para psicólogos', 0),
        ('profile.html: Gestión de perfil y privacidad del usuario', 0),
        ('parent-reports.html: Generación de informes para padres', 0),
        ('videos.html: Videos guiados de respiración y meditación', 0),
        ('active-breaks.html: Configuración de pausas activas', 0),
        ('mental-garden.html: Jardín mental interactivo', 0),
        ('', 0),
        ('CAPA CONTROLLER (LÓGICA DE NEGOCIO)', 3),
        ('', 0),
        ('app.js: Controlador principal de la aplicación, gestión de estado y navegación', 0),
        ('evaluation.js: Lógica de evaluación psicológica, cálculo de puntajes y niveles de riesgo', 0),
        ('chat.js: Gestión de conversaciones con IA, análisis de sentimientos', 0),
        ('exercises.js: Control de ejercicios, temporizadores y seguimiento de progreso', 0),
        ('games.js: Sistema de gamificación, puntos, niveles y logros', 0),
        ('mental-garden.js: Lógica del jardín mental, sistema de plantas y riego', 0),
        ('alerts.js: Gestión de alertas de riesgo alto para psicólogos', 0),
        ('alerts-low-medium.js: Gestión de alertas de riesgo bajo/medio para psicólogos', 0),
        ('hub-client.js: Cliente API para comunicación con backend centralizado', 0),
        ('community.js: Lógica de comunidad, posts, comentarios y moderación', 0),
        ('parent-reports.js: Generación de informes PDF y exportación de datos', 0),
        ('videos.js: Control de reproducción de videos y seguimiento de progreso', 0),
        ('active-breaks.js: Lógica de recordatorios y notificaciones de pausas', 0),
        ('profile.js: Gestión de datos de perfil, privacidad y configuración', 0),
        ('', 0),
        ('ESTRUCTURA DE CARPETAS', 3),
        ('', 0),
        ('web/: Carpeta principal del proyecto web', 0),
        ('web/css/: Archivos de estilos CSS (styles.css, responsive.css)', 0),
        ('web/js/: Archivos JavaScript (app.js, evaluation.js, chat.js, etc.)', 0),
        ('web/assets/: Recursos estáticos (imágenes, iconos, fuentes)', 0),
        ('web/lib/: Librerías externas (frameworks, utilidades)', 0),
        ('web/index.html: Página principal de entrada', 0),
        ('web/package.json: Configuración de dependencias y scripts', 0),
        ('web/README.md: Documentación del proyecto', 0),
    ]
    
    print("2. Insertando nueva arquitectura web correcta...")
    
    # Insertar nueva arquitectura
    for i, (text, level) in enumerate(new_architecture):
        if insert_position + i < len(doc.paragraphs):
            if level == 1:
                new_heading = doc.add_heading(text, level=1)
                remove_spacing(new_heading)
                # Mover al lugar correcto
                ref_paragraph = doc.paragraphs[insert_position + i]
                ref_paragraph._element.addprevious(new_heading._element)
                doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            elif level == 2:
                new_heading = doc.add_heading(text, level=2)
                remove_spacing(new_heading)
                ref_paragraph = doc.paragraphs[insert_position + i]
                ref_paragraph._element.addprevious(new_heading._element)
                doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            elif level == 3:
                new_heading = doc.add_heading(text, level=3)
                remove_spacing(new_heading)
                ref_paragraph = doc.paragraphs[insert_position + i]
                ref_paragraph._element.addprevious(new_heading._element)
                doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            elif text == '':
                new_p = doc.add_paragraph()
                remove_spacing(new_p)
                ref_paragraph = doc.paragraphs[insert_position + i]
                ref_paragraph._element.addprevious(new_p._element)
                doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            else:
                new_p = doc.add_paragraph(text, style='List Bullet')
                remove_spacing(new_p)
                ref_paragraph = doc.paragraphs[insert_position + i]
                ref_paragraph._element.addprevious(new_p._element)
                doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
    
    doc.save(output_doc)
    
    print("3. Verificando estructura corregida...")
    doc = Document(output_doc)
    
    print("   Estructura en el documento:")
    for i, para in enumerate(doc.paragraphs[:40]):
        if para.text.strip():
            print(f"   {i+1}. {para.text[:80]}")
    
    print(f"\n=== DOCUMENTO CORREGIDO CREADO ===")
    print(f"Archivo: {output_doc}")
    print("- Arquitectura técnica corregida con estructura web")
    print("- Archivos HTML, CSS y JavaScript correctos")
    print("- Capas Model, View y Controller web")
    print("- Estructura de carpetas del proyecto web")
    
    return output_doc

if __name__ == "__main__":
    correct_web_structure()
