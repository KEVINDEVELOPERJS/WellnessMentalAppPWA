from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import os
import shutil

def remove_spacing(paragraph):
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def add_section_with_image(doc, title, description, implementation, image_path):
    heading = doc.add_heading(title, level=3)
    remove_spacing(heading)
    
    p = doc.add_paragraph(description)
    remove_spacing(p)
    
    p = doc.add_paragraph(f"Implementado en: {implementation}")
    p.runs[0].font.color.rgb = RGBColor(0, 0, 128)
    remove_spacing(p)
    
    p = doc.add_paragraph("✅ FUNCIONAL")
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    remove_spacing(p)
    
    p = doc.add_paragraph("Evidencia Visual:")
    p.runs[0].font.italic = True
    remove_spacing(p)
    
    doc.add_picture(image_path, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(last_paragraph)
    
    doc.add_paragraph()

def create_final_document():
    original_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp\INFORME_FINAL_INTEGRADO.docx'
    web_images_doc = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.docx'
    output_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_FINAL_COMPLETO.docx'
    
    print("CREANDO DOCUMENTO FINAL COMPLETO")
    print("1. Copiando documento original...")
    shutil.copy2(original_doc, output_doc)
    
    replacements = {
        'Android': 'Web',
        'android': 'web',
        'APK': 'Web',
        'apk': 'web',
        'Android Studio': 'HTML5/CSS3/JavaScript',
        'Kotlin': 'JavaScript',
        'Java': 'JavaScript',
        'XML': 'HTML',
        'Gradle': 'NPM',
        'Room Database': 'IndexedDB',
        'SQLite': 'IndexedDB',
        'SharedPreferences': 'LocalStorage',
        'Retrofit': 'Fetch API',
        'Activity': 'Page',
        'aplicación móvil': 'aplicación web',
        'app móvil': 'app web',
        'móvil': 'web',
        'native': 'web',
        'nativo': 'web',
    }
    
    print("2. Modificando texto...")
    doc = Document(output_doc)
    
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)
    
    for paragraph in doc.paragraphs:
        if 'Android' in paragraph.text and ('INFORME' in paragraph.text or 'Funcionamiento' in paragraph.text):
            paragraph.text = paragraph.text.replace('Android', 'Web')
            if 'Wellness Mental App' in paragraph.text:
                paragraph.text = paragraph.text.replace('Wellness Mental App', 'Wellness Mental App Web')
    
    doc.save(output_doc)
    
    print("3. Extrayendo imágenes web...")
    web_images = {}
    with zipfile.ZipFile(web_images_doc, 'r') as z:
        for f in z.namelist():
            if f.startswith('word/media/image'):
                num = f.replace('word/media/image', '').replace('.png', '')
                web_images[num] = z.read(f)
    
    print(f"   Imágenes: {len(web_images)}")
    
    print("4. Reemplazando imágenes...")
    temp = output_doc + '.tmp'
    
    with zipfile.ZipFile(output_doc, 'r') as z1:
        with zipfile.ZipFile(temp, 'w', zipfile.ZIP_DEFLATED) as z2:
            replaced = 0
            for info in z1.infolist():
                data = z1.read(info.filename)
                if info.filename.startswith('word/media/image'):
                    num = info.filename.replace('word/media/image', '').replace('.png', '')
                    if num in web_images:
                        z2.writestr(info.filename, web_images[num])
                        replaced += 1
                    else:
                        z2.writestr(info, data)
                else:
                    z2.writestr(info, data)
            print(f"   Reemplazadas: {replaced}")
    
    os.remove(output_doc)
    shutil.move(temp, output_doc)
    
    print("5. Añadiendo módulos adicionales...")
    doc = Document(output_doc)
    
    doc.add_page_break()
    
    heading = doc.add_heading('Módulos Adicionales del Sistema Web', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(heading)
    
    heading2 = doc.add_heading('Criterios de Aceptación Implementados', level=2)
    remove_spacing(heading2)
    
    additional_modules = {
        '35': {'title': 'CA-Extra.1: Vista Web General del Sistema', 'description': 'Vista general del sistema web mostrando la integración de todos los módulos.', 'implementation': 'index.html, app.js'},
        '36': {'title': 'CA-Extra.2: Vista Web de Navegación', 'description': 'Sistema de navegación web entre diferentes módulos.', 'implementation': 'app.js'},
        '37': {'title': 'CA-Extra.3: Vista Web de Perfil de Usuario', 'description': 'Vista del perfil de usuario web con estadísticas.', 'implementation': 'profile.html, profile.js'},
        '38': {'title': 'CA-Extra.4: Vista Web de Configuración', 'description': 'Panel de configuración web con opciones personalización.', 'implementation': 'profile.html, profile.js'},
        '39': {'title': 'CA-02.1: Dashboard Web de Psicólogo - Vista General', 'description': 'Panel principal del psicólogo con vista general de estudiantes y alertas.', 'implementation': 'index.html, app.js'},
        '40': {'title': 'CA-02.1: Dashboard Web de Psicólogo - Panel de Alertas', 'description': 'Panel específico de alertas con casos clasificados por riesgo.', 'implementation': 'index.html, app.js'},
        '41': {'title': 'CA-Alerta.1: Detalle Web de Alerta de Riesgo Alto', 'description': 'Vista detallada de alerta de riesgo alto con información del estudiante.', 'implementation': 'alerts.html, alerts.js'},
        '42': {'title': 'CA-Alerta.2: Panel Web de Alertas de Riesgo Bajo/Medio', 'description': 'Panel de alertas para seguimiento preventivo de riesgo bajo/medio.', 'implementation': 'alerts-low-medium.html, alerts-low-medium.js'},
        '43': {'title': 'CA-Extra.5: Vista Web de Reportes y Estadísticas', 'description': 'Panel de reportes web con estadísticas generales del sistema.', 'implementation': 'parent-reports.html, parent-reports.js'},
        '44': {'title': 'CA-Extra.6: Vista Web de Configuración del Sistema', 'description': 'Vista de configuración general del sistema web administrativo.', 'implementation': 'app.js'},
    }
    
    for num in range(35, 45):
        num_str = str(num)
        if num_str in web_images and num_str in additional_modules:
            module = additional_modules[num_str]
            temp_img = f'temp_{num_str}.png'
            with open(temp_img, 'wb') as f:
                f.write(web_images[num_str])
            
            try:
                add_section_with_image(doc, module['title'], module['description'], module['implementation'], temp_img)
                print(f"   Módulo {num_str}: {module['title']}")
            except Exception as e:
                print(f"   Error {num_str}: {e}")
            finally:
                if os.path.exists(temp_img):
                    os.remove(temp_img)
    
    doc.save(output_doc)
    
    print("6. Ajustando tamaño imágenes...")
    doc = Document(output_doc)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for shape in run._element.xpath('.//pic:pic'):
                for extent in shape.xpath('.//a:ext'):
                    extent.set('{http://schemas.openxmlformats.org/drawingml/2006/main}cx', str(int(Inches(6.5).emu)))
                    extent.set('{http://schemas.openxmlformats.org/drawingml/2006/main}cy', str(int(Inches(4.5).emu)))
    
    doc.save(output_doc)
    
    print(f"\nDOCUMENTO FINAL: {output_doc}")
    print(f"- Texto modificado")
    print(f"- {replaced} imágenes reemplazadas")
    print(f"- {len(additional_modules)} módulos adicionales")
    print(f"- Dashboard psicólogo en imagen 39")
    print(f"- Tamaño 6.5\" x 4.5\"")
    
    return output_doc

if __name__ == "__main__":
    create_final_document()
