from docx import Document
from docx.shared import Inches
import zipfile
import os
import shutil

def modify_document_to_web():
    """Modifica el documento copia exacta para contenido web"""
    
    input_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_COPIA_EXACTA.docx'
    web_images_doc = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.docx'
    output_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_FINAL.docx'
    
    print("1. Copiando documento base...")
    shutil.copy2(input_doc, output_doc)
    
    # Reemplazos básicos de Android a Web
    replacements = {
        'Android': 'Web',
        'android': 'web',
        'APK': 'Web',
        'apk': 'web',
        'Android Studio': 'HTML5/CSS3/JavaScript',
        'Kotlin': 'JavaScript',
        'Java': 'JavaScript',
        'XML': 'HTML',
        'Gradle': 'NPM',
        'Room Database': 'IndexedDB',
        'SQLite': 'IndexedDB',
        'SharedPreferences': 'LocalStorage',
        'Retrofit': 'Fetch API',
        'actividad': 'página',
        'Activity': 'Page',
        'aplicación móvil': 'aplicación web',
        'app móvil': 'app web',
        'móvil': 'web',
        'native': 'web',
        'nativo': 'web',
        'Studio': 'VS Code',
        'Emulator': 'Browser',
        'Manifest': 'Configuración',
        'build.gradle': 'package.json',
        'styles.xml': 'CSS',
        'res/': 'assets/',
        '.kt': '.js',
        '.java': '.js',
        '.xml': '.html',
    }
    
    print("2. Modificando texto de Android a Web...")
    doc = Document(output_doc)
    
    # Modificar parrafos
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        for old, new in replacements.items():
            if old in original_text:
                paragraph.text = original_text.replace(old, new)
    
    # Modificar tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    for old, new in replacements.items():
                        if old in original_text:
                            paragraph.text = original_text.replace(old, new)
    
    # Modificar título principal específicamente
    for paragraph in doc.paragraphs:
        if 'Android' in paragraph.text and ('INFORME' in paragraph.text or 'Funcionamiento' in paragraph.text):
            paragraph.text = paragraph.text.replace('Android', 'Web')
            if 'Wellness Mental App' in paragraph.text:
                paragraph.text = paragraph.text.replace('Wellness Mental App', 'Wellness Mental App Web')
    
    print("3. Ajustando tamaño de imágenes...")
    # Ajustar imágenes a 6 pulgadas para mejor visibilidad
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for shape in run._element.xpath('.//pic:pic'):
                # Encontrar el elemento de imagen y ajustar tamaño
                for blip in shape.xpath('.//a:blip'):
                    # Buscar el elemento de tamaño
                    for extent in shape.xpath('.//a:ext'):
                        # Ajustar el tamaño a 6 pulgadas (approx 5.7M en EMUs)
                        extent.set('{http://schemas.openxmlformats.org/drawingml/2006/main}cx', str(int(Inches(6).emu)))
                        extent.set('{http://schemas.openxmlformats.org/drawingml/2006/main}cy', str(int(Inches(4).emu)))
    
    doc.save(output_doc)
    
    print("4. Reinsertando todas las imágenes...")
    # Extraer todas las imágenes del documento web
    web_images = {}
    with zipfile.ZipFile(web_images_doc, 'r') as z:
        for f in z.namelist():
            if f.startswith('word/media/image'):
                num = f.replace('word/media/image', '').replace('.png', '')
                web_images[num] = z.read(f)
    
    print(f"   Imágenes web disponibles: {len(web_images)}")
    
    # Reemplazar imágenes
    temp = output_doc + '.tmp'
    with zipfile.ZipFile(output_doc, 'r') as z1:
        with zipfile.ZipFile(temp, 'w', zipfile.ZIP_DEFLATED) as z2:
            replaced_count = 0
            for info in z1.infolist():
                data = z1.read(info.filename)
                if info.filename.startswith('word/media/image'):
                    num = info.filename.replace('word/media/image', '').replace('.png', '')
                    if num in web_images:
                        z2.writestr(info.filename, web_images[num])
                        replaced_count += 1
                    else:
                        z2.writestr(info, data)
                else:
                    z2.writestr(info, data)
            print(f"   Imágenes reemplazadas: {replaced_count}")
    
    os.remove(output_doc)
    shutil.move(temp, output_doc)
    
    print(f"5. Documento final creado: {output_doc}")
    print("   - Texto modificado de Android a Web")
    print(f"   - {len(web_images)} imágenes incluidas")
    print("   - Tamaño de imágenes ajustado para mejor visibilidad")
    print("   - Estructura original mantenida")
    
    return output_doc

if __name__ == "__main__":
    modify_document_to_web()
