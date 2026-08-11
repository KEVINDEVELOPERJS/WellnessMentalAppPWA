from docx import Document
from docx.shared import Inches
import zipfile
import os
import shutil

def replace_all_images():
    """Reemplaza todas las imágenes con las del documento web"""
    
    original_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp\INFORME_FINAL_INTEGRADO.docx'
    web_images_doc = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.docx'
    output_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_COMPLETO.docx'
    
    print("1. Copiando documento original...")
    shutil.copy2(original_doc, output_doc)
    
    # Modificar texto
    replacements = {
        'Android': 'Web',
        'android': 'web',
        'APK': 'Web',
        'apk': 'web',
        'Android Studio': 'HTML5/CSS3/JavaScript',
        'Kotlin': 'JavaScript',
        'Java': 'JavaScript',
        'XML': 'HTML',
        'Gradle': 'NPM',
        'Room Database': 'IndexedDB',
        'SQLite': 'IndexedDB',
        'SharedPreferences': 'LocalStorage',
        'Retrofit': 'Fetch API',
        'Activity': 'Page',
        'aplicación móvil': 'aplicación web',
        'app móvil': 'app web',
        'móvil': 'web',
        'native': 'web',
        'nativo': 'web',
    }
    
    print("2. Modificando texto...")
    doc = Document(output_doc)
    
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)
    
    # Modificar título
    for paragraph in doc.paragraphs:
        if 'Android' in paragraph.text and ('INFORME' in paragraph.text or 'Funcionamiento' in paragraph.text):
            paragraph.text = paragraph.text.replace('Android', 'Web')
            if 'Wellness Mental App' in paragraph.text:
                paragraph.text = paragraph.text.replace('Wellness Mental App', 'Wellness Mental App Web')
    
    doc.save(output_doc)
    
    print("3. Reemplazando imágenes...")
    web_images = {}
    with zipfile.ZipFile(web_images_doc, 'r') as z:
        for f in z.namelist():
            if f.startswith('word/media/image'):
                num = f.replace('word/media/image', '').replace('.png', '')
                web_images[num] = z.read(f)
    
    print(f"   Imágenes web: {len(web_images)}")
    
    temp = output_doc + '.tmp'
    with zipfile.ZipFile(output_doc, 'r') as z1:
        with zipfile.ZipFile(temp, 'w', zipfile.ZIP_DEFLATED) as z2:
            replaced = 0
            for info in z1.infolist():
                data = z1.read(info.filename)
                if info.filename.startswith('word/media/image'):
                    num = info.filename.replace('word/media/image', '').replace('.png', '')
                    if num in web_images:
                        z2.writestr(info.filename, web_images[num])
                        replaced += 1
                    else:
                        z2.writestr(info, data)
                else:
                    z2.writestr(info, data)
            print(f"   Imágenes reemplazadas: {replaced}")
    
    os.remove(output_doc)
    shutil.move(temp, output_doc)
    
    print(f"4. Documento final: {output_doc}")
    print("   - Texto modificado")
    print(f"   - {replaced} imágenes reemplazadas")
    
    return output_doc

if __name__ == "__main__":
    replace_all_images()
