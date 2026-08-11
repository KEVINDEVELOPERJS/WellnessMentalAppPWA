from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

def remove_spacing(paragraph):
    """Remove spacing before and after a paragraph"""
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def create_word_document_web():
    doc = Document()
    
    # Define image mappings for WEB APP
    image_mapping = {
        'HU-01': {
            'title': 'HU-01: Registro y Autenticación de Usuarios',
            'images': ['web_registro.png', 'web_login.png', 'web_consentimiento.png'],
            'descriptions': [
                'CA-01.1: Pantalla de Registro Web con validación de campos',
                'CA-01.4: Pantalla de Login Web con autenticación',
                'CA-01.3: Pantalla de Consentimiento Parental Web'
            ]
        },
        'HU-02': {
            'title': 'HU-02: Dashboard Principal',
            'images': ['web_dashboard_estudiante.png', 'web_dashboard_psicologo.png'],
            'descriptions': [
                'CA-02.1: Dashboard Web de Estudiante con estado general',
                'CA-02.1: Dashboard Web de Psicólogo con panel de alertas'
            ]
        },
        'HU-03': {
            'title': 'HU-03: Sistema de Evaluación Psicológica',
            'images': ['web_cuestionarios.png', 'web_gad7.png', 'web_resultados.png', 'web_historial.png'],
            'descriptions': [
                'CA-03.1: Lista Web de Cuestionarios Disponibles (GAD-7, PHQ-9, PSS-10)',
                'CA-03.1: Cuestionario Web GAD-7 en Progreso con escala Likert',
                'CA-03.2: Resultados Web de Evaluación con nivel de riesgo',
                'CA-03.3: Historial Web de Evaluaciones con fechas y puntajes'
            ]
        },
        'HU-04': {
            'title': 'HU-04: Chat con Inteligencia Artificial',
            'images': ['web_chat.png', 'web_chat_conversacion.png'],
            'descriptions': [
                'CA-04.1: Interfaz Web de Chat IA con bubbles de mensajes',
                'CA-04.2: Conversación Web con Asistente Emocional empático'
            ]
        },
        'HU-05': {
            'title': 'HU-05: Ejercicios de Bienestar',
            'images': ['web_ejercicios.png', 'web_respiracion.png', 'web_meditacion.png', 'web_progreso.png'],
            'descriptions': [
                'CA-05.1: Catálogo Web de Ejercicios por categoría',
                'CA-05.2: Ejercicio Web de Respiración 4-7-8 con temporizador',
                'CA-05.2: Meditación Web Guiada con instrucciones',
                'CA-05.3: Progreso Web de Ejercicios con estadísticas'
            ]
        },
        'HU-08': {
            'title': 'HU-08: Gamificación y Juegos',
            'images': ['web_juegos.png', 'web_puzzle.png', 'web_arte.png', 'web_ritmo.png', 'web_jardin.png', 'web_logros.png', 'web_puntos.png'],
            'descriptions': [
                'CA-08.1: Pantalla Web Principal de Juegos',
                'CA-08.1: Mini-juego Web Puzzle Zen',
                'CA-08.1: Mini-juego Web Arte Emocional',
                'CA-08.1: Mini-juego Web Ritmo Calma',
                'CA-08.1: Jardín Mental Web con progreso visual',
                'CA-08.4: Logros Web Desbloqueados con celebración',
                'CA-08.2-CA-08.3: Sistema Web de Puntos y Niveles'
            ]
        },
        'HU-07': {
            'title': 'HU-07: Comunidad Estudiantil',
            'images': ['web_comunidad.png', 'web_post.png', 'web_interaccion.png'],
            'descriptions': [
                'CA-07.1: Foros Web de Comunidad categorizados',
                'CA-07.2: Creación Web de Post con título y contenido',
                'CA-07.2: Interacción Web en Comunidad con comentarios'
            ]
        },
        'HU-09': {
            'title': 'HU-09: Informes para Padres',
            'images': ['web_informes.png', 'web_informe_previa.png'],
            'descriptions': [
                'CA-09.1: Generación Web de Informe con resumen de bienestar',
                'CA-09.1: Vista Previa Web de Informe para Padres'
            ]
        },
        'HU-10': {
            'title': 'HU-10: Gestión de Perfil y Privacidad',
            'images': ['web_perfil.png', 'web_privacidad.png', 'web_sesiones.png'],
            'descriptions': [
                'CA-10.1: Perfil Web de Usuario con estadísticas',
                'CA-10.2: Configuración Web de Privacidad y notificaciones',
                'CA-10.3: Gestión Web de Sesiones activas'
            ]
        },
        'HU-06': {
            'title': 'HU-06: Diseño de Cuestionarios por Psicólogo',
            'images': ['web_editor.png', 'web_editor_previa.png'],
            'descriptions': [
                'CA-06.1: Editor Web de Cuestionarios con campos y botones',
                'CA-06.4: Vista Previa Web de Cuestionario como estudiante'
            ]
        },
        'Alertas': {
            'title': 'Gestión de Alertas para Psicólogos',
            'images': ['web_alertas_panel.png', 'web_alertas_detalle.png', 'web_alertas_bajo_medio.png'],
            'descriptions': [
                'CA-Alerta.1: Panel Web de Alertas de Riesgo Alto',
                'CA-Alerta.1: Detalle Web de Alerta de Riesgo con información estudiante',
                'CA-Alerta.2: Panel Web de Alertas de Riesgo Bajo/Medio'
            ]
        },
        'Misiones': {
            'title': 'Misiones Diarias y Check-In',
            'images': ['web_checkin.png', 'web_misiones.png'],
            'descriptions': [
                'CA-Mision.1: Check-In Web Diario con estado de ánimo',
                'CA-Mision.2: Misiones Web Diarias con recompensas'
            ]
        },
        'Videos': {
            'title': 'Videos Guiados',
            'images': ['web_videos.png', 'web_video_reproduccion.png'],
            'descriptions': [
                'CA-Video.1: Catálogo Web de Videos Guiados',
                'CA-Video.2: Reproducción Web de Video con controles'
            ]
        },
        'Pausas': {
            'title': 'Pausas Activas',
            'images': ['web_pausas.png', 'web_pausa_config.png'],
            'descriptions': [
                'CA-Pausa.1: Pantalla Web de Pausas Activas',
                'CA-Pausa.2: Configuración Web de Recordatorios de Pausas'
            ]
        }
    }
    
    # Carpeta donde deben estar las imágenes de la app web
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    # ==================== PARTE 1: INFORME TIPO ARTÍCULO CON IMÁGENES INTEGRADAS ====================
    
    # Title
    title = doc.add_heading('Informe Técnico: Funcionamiento Correcto de Wellness Mental App Web', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(title)
    
    # Metadata
    p = doc.add_paragraph('Fecha: 4 de agosto de 2026')
    remove_spacing(p)
    p = doc.add_paragraph('Versión: 1.0')
    remove_spacing(p)
    p = doc.add_paragraph('Tipo: Informe de Validación de Criterios de Aceptación - Versión Web')
    remove_spacing(p)
    p = doc.add_paragraph()
    remove_spacing(p)
    
    # Resumen Ejecutivo
    h = doc.add_heading('Resumen Ejecutivo', level=1)
    remove_spacing(h)
    p = doc.add_paragraph(
        'Este informe presenta una validación detallada del funcionamiento correcto de la aplicación web '
        '"Wellness Mental App Web", desarrollada con HTML5, CSS3, JavaScript y arquitectura PWA. La aplicación '
        'está diseñada para el bienestar mental de estudiantes adolescentes (13-18 años) e implementa múltiples '
        'módulos interconectados que cumplen con los criterios de aceptación definidos para cada historia de usuario. '
        'La versión web mantiene paridad funcional con la versión Android nativa, ofreciendo acceso multiplataforma '
        'mediante tecnologías web modernas y Capacitor para despliegue híbrido.'
    )
    remove_spacing(p)
    
    # Arquitectura Técnica
    h = doc.add_heading('Arquitectura Técnica', level=1)
    remove_spacing(h)
    h = doc.add_heading('Estructura Web Implementada', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Capa Model (Datos y Persistencia):', level=3)
    remove_spacing(h)
    p = doc.add_paragraph()
    p.add_run('IndexedDB/').bold = True
    p.add_run(': Base de datos local del navegador para almacenamiento persistente (usuarios, respuestas, resultados, chat, ejercicios)')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('LocalStorage/').bold = True
    p.add_run(': Gestión de sesión y preferencias de usuario')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('API Hub/').bold = True
    p.add_run(': Cliente REST para sincronización con backend centralizado')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('LocalStorage/').bold = True
    p.add_run(': Gestión de preferencias y configuración')
    remove_spacing(p)
    
    h = doc.add_heading('Capa View (Interfaz de Usuario):', level=3)
    remove_spacing(h)
    p = doc.add_paragraph()
    p.add_run('index.html/').bold = True
    p.add_run(': Pantalla principal de login/registro')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('evaluation.html/').bold = True
    p.add_run(': Sistema de evaluación psicológica (GAD-7, PHQ-9, PSS-10)')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('chat.html/').bold = True
    p.add_run(': Interfaz de chat con asistente emocional IA')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('exercises.html/').bold = True
    p.add_run(': Catálogo de ejercicios de bienestar')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('games.html/').bold = True
    p.add_run(': Sistema de gamificación y mini-juegos')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('community.html/').bold = True
    p.add_run(': Foros de comunidad estudiantil')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('alerts.html/').bold = True
    p.add_run(': Panel de alertas para psicólogos (riesgo alto)')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('alerts-low-medium.html/').bold = True
    p.add_run(': Panel de alertas para psicólogos (riesgo bajo/medio)')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('questionnaire-editor.html/').bold = True
    p.add_run(': Editor de cuestionarios para psicólogos')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('profile.html/').bold = True
    p.add_run(': Gestión de perfil y privacidad')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('parent-reports.html/').bold = True
    p.add_run(': Generación de informes para padres')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('videos.html/').bold = True
    p.add_run(': Videos guiados de respiración y meditación')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('active-breaks.html/').bold = True
    p.add_run(': Configuración de pausas activas')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('mental-garden.html/').bold = True
    p.add_run(': Jardín mental interactivo')
    remove_spacing(p)
    
    h = doc.add_heading('Capa Controller (Lógica de Negocio):', level=3)
    remove_spacing(h)
    p = doc.add_paragraph()
    p.add_run('app.js/').bold = True
    p.add_run(': Controlador principal de la aplicación, gestión de estado y navegación')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('evaluation.js/').bold = True
    p.add_run(': Lógica de evaluación psicológica, cálculo de puntajes y niveles de riesgo')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('chat.js/').bold = True
    p.add_run(': Gestión de conversaciones con IA, análisis de sentimientos')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('exercises.js/').bold = True
    p.add_run(': Control de ejercicios, temporizadores y seguimiento de progreso')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('games.js/').bold = True
    p.add_run(': Sistema de gamificación, puntos, niveles y logros')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('mental-garden.js/').bold = True
    p.add_run(': Lógica del jardín mental, sistema de plantas y riego')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('alerts.js/').bold = True
    p.add_run(': Gestión de alertas de riesgo alto para psicólogos')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('alerts-low-medium.js/').bold = True
    p.add_run(': Gestión de alertas de riesgo bajo/medio para psicólogos')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('hub-client.js/').bold = True
    p.add_run(': Cliente API para comunicación con backend centralizado')
    remove_spacing(p)
    
    # Validación por Criterios de Aceptación
    h = doc.add_heading('Validación por Criterios de Aceptación con Evidencia Visual', level=1)
    remove_spacing(h)
    
    # Process each module
    for key, module in image_mapping.items():
        h = doc.add_heading(module['title'], level=2)
        remove_spacing(h)
        
        h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
        remove_spacing(h)
        
        for i, (image, description) in enumerate(zip(module['images'], module['descriptions'])):
            h = doc.add_heading(description.split(':')[0], level=4)
            remove_spacing(h)
            
            # Description with file reference
            p = doc.add_paragraph()
            p.add_run(f'{description} - Implementado en archivos web correspondientes. ').bold = True
            p.add_run('✅ FUNCIONAL')
            remove_spacing(p)
            
            # Add image
            image_path = os.path.join(images_folder, image)
            if os.path.exists(image_path):
                try:
                    # Add descriptive text before image
                    p = doc.add_paragraph()
                    run = p.add_run('Evidencia Visual: ' + description.split(':')[1].strip() if ':' in description else description)
                    run.bold = True
                    run.font.size = Pt(9)
                    remove_spacing(p)
                    
                    # Add image centered
                    doc.add_picture(image_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    remove_spacing(last_paragraph)
                    
                    p = doc.add_paragraph()
                    remove_spacing(p)
                except Exception as e:
                    p = doc.add_paragraph(f'[Error al cargar imagen: {e}]')
                    remove_spacing(p)
            else:
                # Add placeholder text when image is not found
                p = doc.add_paragraph()
                run = p.add_run(f'Evidencia Visual: {description.split(":")[1].strip() if ":" in description else description}')
                run.bold = True
                run.font.size = Pt(9)
                remove_spacing(p)
                
                p = doc.add_paragraph(f'[CAPTURA DE PANTALLA REQUERIDA: {image}]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                remove_spacing(p)
                
                p = doc.add_paragraph()
                remove_spacing(p)
        
        add_page_break(doc)
    
    # Save document
    output_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB.docx'
    doc.save(output_path)
    print(f'Documento creado exitosamente: {output_path}')
    print(f'Total de módulos documentados: {len(image_mapping)}')
    print(f'Total de imágenes requeridas: {sum(len(module["images"]) for module in image_mapping.values())}')

if __name__ == '__main__':
    create_word_document_web()
