from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os
import glob

def remove_spacing(paragraph):
    """Remove spacing before and after a paragraph"""
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def create_web_document_with_images():
    """Crea el documento Word basado en el contenido web existente con imágenes"""
    
    # Leer el contenido del informe web
    content_file = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_WEB_COMPLETO.txt'
    
    with open(content_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Carpeta de imágenes
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    # Obtener lista de imágenes disponibles
    available_images = []
    if os.path.exists(images_folder):
        for ext in ['*.png', '*.jpg', '*.jpeg', '*.gif']:
            available_images.extend(glob.glob(os.path.join(images_folder, ext)))
        available_images.sort()
    
    print(f"Imágenes disponibles: {len(available_images)}")
    for img in available_images:
        print(f"  - {os.path.basename(img)}")
    
    # Mapeo de imágenes para cada sección
    image_mapping = {
        'HU-01': ['web_registro.png', 'web_login.png', 'web_consentimiento.png'],
        'HU-02': ['web_dashboard_estudiante.png', 'web_dashboard_psicologo.png'],
        'HU-03': ['web_cuestionarios.png', 'web_gad7.png', 'web_resultados.png', 'web_historial.png'],
        'HU-04': ['web_chat.png', 'web_chat_conversacion.png'],
        'HU-05': ['web_ejercicios.png', 'web_respiracion.png', 'web_meditacion.png', 'web_progreso.png'],
        'HU-08': ['web_juegos.png', 'web_puzzle.png', 'web_arte.png', 'web_ritmo.png', 'web_jardin.png', 'web_logros.png', 'web_puntos.png'],
        'HU-07': ['web_comunidad.png', 'web_post.png', 'web_interaccion.png'],
        'HU-09': ['web_informes.png', 'web_informe_previa.png'],
        'HU-10': ['web_perfil.png', 'web_privacidad.png', 'web_sesiones.png'],
        'HU-06': ['web_editor.png', 'web_editor_previa.png'],
        'Alertas': ['web_alertas_panel.png', 'web_alertas_detalle.png', 'web_alertas_bajo_medio.png'],
        'Misiones': ['web_checkin.png', 'web_misiones.png'],
        'Videos': ['web_videos.png', 'web_video_reproduccion.png'],
        'Pausas': ['web_pausas.png', 'web_pausa_config.png']
    }
    
    current_section = None
    image_index = 0
    
    # Procesar el contenido línea por línea
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # Detectar sección actual
        if line.startswith('HU-'):
            current_section = line.split(':')[0]  # HU-01, HU-02, etc.
            image_index = 0
        elif 'Alertas' in line:
            current_section = 'Alertas'
            image_index = 0
        elif 'Misiones' in line:
            current_section = 'Misiones'
            image_index = 0
        elif 'Videos' in line:
            current_section = 'Videos'
            image_index = 0
        elif 'Pausas' in line:
            current_section = 'Pausas'
            image_index = 0
        
        # Títulos principales (con ========)
        if line.startswith('=') and line.endswith('='):
            title_text = line.replace('=', '').strip()
            heading = doc.add_heading(title_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            remove_spacing(heading)
        
        # Subtítulos (con --------)
        elif line.startswith('-') and line.endswith('-'):
            subtitle_text = line.replace('-', '').strip()
            heading = doc.add_heading(subtitle_text, level=2)
            remove_spacing(heading)
        
        # Líneas que parecen títulos de módulos (HU-XX: ...)
        elif line.startswith('HU-') or line.startswith('CA-'):
            if line.startswith('HU-'):
                heading = doc.add_heading(line, level=2)
            else:
                heading = doc.add_heading(line, level=3)
            remove_spacing(heading)
        
        # Líneas de evidencia visual - insertar imagen
        elif 'Evidencia Visual:' in line or 'Captura:' in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.italic = True
            remove_spacing(p)
            
            # Intentar insertar imagen correspondiente
            if current_section and current_section in image_mapping:
                if image_index < len(image_mapping[current_section]):
                    image_name = image_mapping[current_section][image_index]
                    image_path = os.path.join(images_folder, image_name)
                    
                    if os.path.exists(image_path):
                        try:
                            doc.add_picture(image_path, width=Inches(5))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            remove_spacing(last_paragraph)
                            print(f"Imagen insertada: {image_name}")
                            image_index += 1
                        except Exception as e:
                            print(f"Error insertando imagen {image_name}: {e}")
                    else:
                        print(f"Imagen no encontrada: {image_path}")
                        # Usar imagen disponible si existe
                        if available_images:
                            try:
                                doc.add_picture(available_images[0], width=Inches(5))
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                remove_spacing(last_paragraph)
                                print(f"Usando imagen alternativa: {os.path.basename(available_images[0])}")
                                image_index += 1
                            except Exception as e:
                                print(f"Error con imagen alternativa: {e}")
        
        # Líneas con indicadores de funcionalidad (✅)
        elif '✅' in line or 'FUNCIONAL' in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Verde
            remove_spacing(p)
        
        # Líneas de implementación
        elif line.startswith('Implementado en:'):
            p = doc.add_paragraph(line)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Azul
            remove_spacing(p)
        
        # Texto normal
        else:
            p = doc.add_paragraph(line)
            remove_spacing(p)
    
    # Guardar el documento
    output_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_CON_IMAGENES.docx'
    doc.save(output_path)
    print(f"Documento guardado exitosamente en: {output_path}")
    return output_path

if __name__ == "__main__":
    create_web_document_with_images()
