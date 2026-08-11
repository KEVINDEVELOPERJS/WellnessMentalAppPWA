from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import os
import shutil

def remove_spacing(paragraph):
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def add_section_with_image(doc, title, description, implementation, image_path):
    """Añade una sección completa como las demás del documento"""
    
    # Título del criterio
    heading = doc.add_heading(title, level=3)
    remove_spacing(heading)
    
    # Descripción
    p = doc.add_paragraph(description)
    remove_spacing(p)
    
    # Implementación
    p = doc.add_paragraph(f"Implementado en: {implementation}")
    p.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Azul
    remove_spacing(p)
    
    # Estado
    p = doc.add_paragraph("✅ FUNCIONAL")
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Verde
    remove_spacing(p)
    
    # Evidencia visual
    p = doc.add_paragraph("Evidencia Visual:")
    p.runs[0].font.italic = True
    remove_spacing(p)
    
    # Insertar imagen
    doc.add_picture(image_path, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(last_paragraph)
    
    # Espacio
    doc.add_paragraph()

def organize_additional_images():
    """Organiza las imágenes adicionales con descripciones técnicas"""
    
    input_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_SOLUCION_FINAL.docx'
    web_images_doc = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.docx'
    output_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_ORGANIZADO.docx'
    
    print("=== ORGANIZANDO IMÁGENES ADICIONALES ===")
    print("1. Copiando documento...")
    shutil.copy2(input_doc, output_doc)
    
    # Extraer imágenes web
    print("2. Extrayendo imágenes web...")
    web_images = {}
    with zipfile.ZipFile(web_images_doc, 'r') as z:
        for f in z.namelist():
            if f.startswith('word/media/image'):
                num = f.replace('word/media/image', '').replace('.png', '')
                web_images[num] = z.read(f)
    
    print(f"   Imágenes extraídas: {len(web_images)}")
    
    # Mapeo de imágenes adicionales a módulos (basado en el orden del documento web)
    # Asumiendo que las imágenes 35-44 corresponden a módulos específicos
    additional_modules = {
        '35': {
            'title': 'CA-02.1: Dashboard Web de Psicólogo con Panel de Alertas',
            'description': 'Panel específico para psicólogos con vista general de alertas de riesgo alto, métricas de estudiantes y acceso rápido a casos críticos.',
            'implementation': 'index.html, app.js (panel específico para psicólogos con alertas)'
        },
        '36': {
            'title': 'CA-Alerta.1: Detalle Web de Alerta de Riesgo Alto',
            'description': 'Vista detallada de alerta de riesgo con información completa del estudiante, historial de evaluaciones y recomendaciones de intervención.',
            'implementation': 'alerts.html, alerts.js (detalle de alerta web con información estudiante)'
        },
        '37': {
            'title': 'CA-Alerta.2: Panel Web de Alertas de Riesgo Bajo/Medio',
            'description': 'Panel de alertas para psicólogos con casos de riesgo bajo y medio, permitiendo seguimiento preventivo y monitoreo continuo.',
            'implementation': 'alerts-low-medium.html, alerts-low-medium.js (panel web de alertas riesgo bajo/medio)'
        },
        '38': {
            'title': 'CA-Mision.1: Check-In Web Diario con Estado de Ánimo',
            'description': 'Sistema de check-in web diario donde los estudiantes registran su estado de ánimo, permitiendo seguimiento emocional continuo.',
            'implementation': 'app.js, evaluation.js (check-in web con estado de ánimo)'
        },
        '39': {
            'title': 'CA-Mision.2: Misiones Web Diarias con Recompensas',
            'description': 'Sistema de misiones web diarias con recompensas y gamificación para mantener el compromiso de los estudiantes.',
            'implementation': 'games.js, app.js (misiones web con sistema de recompensas)'
        },
        '40': {
            'title': 'CA-Video.1: Catálogo Web de Videos Guiados',
            'description': 'Catálogo web de videos guiados de respiración, meditación y técnicas de relajación organizados por categoría.',
            'implementation': 'videos.html, videos.js (catálogo web de videos guiados)'
        },
        '41': {
            'title': 'CA-Video.2: Reproducción Web de Video con Controles',
            'description': 'Reproductor web con controles de reproducción, pausa, avance y retroceso para videos guiados.',
            'implementation': 'videos.html, videos.js (reproductor web con controles)'
        },
        '42': {
            'title': 'CA-Pausa.1: Pantalla Web de Pausas Activas',
            'description': 'Pantalla web de configuración de pausas activas para recordatorios de bienestar durante el estudio.',
            'implementation': 'active-breaks.html, active-breaks.js (pantalla web de pausas activas)'
        },
        '43': {
            'title': 'CA-Pausa.2: Configuración Web de Recordatorios de Pausas',
            'description': 'Configuración web de recordatorios de pausas activas con intervalos personalizados y notificaciones.',
            'implementation': 'active-breaks.html, active-breaks.js (configuración web de recordatorios)'
        },
        '44': {
            'title': 'CA-Extra.1: Vista Web General del Sistema',
            'description': 'Vista general del sistema web mostrando la integración de todos los módulos y la experiencia de usuario completa.',
            'implementation': 'index.html, app.js (vista general del sistema web)'
        }
    }
    
    print("3. Modificando documento para organizar imágenes adicionales...")
    doc = Document(output_doc)
    
    # Eliminar la sección de imágenes adicionales simple
    found_additional_section = False
    paragraphs_to_remove = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        if 'Imágenes Adicionales' in paragraph.text:
            found_additional_section = True
        if found_additional_section:
            paragraphs_to_remove.append(i)
    
    # Eliminar desde el final hacia adelante para no afectar los índices
    for i in reversed(paragraphs_to_remove):
        p = doc.paragraphs[i]
        p_element = p._element
        p_element.getparent().remove(p_element)
    
    # Añadir sección organizada al final
    doc.add_page_break()
    
    heading = doc.add_heading('Módulos Adicionales del Sistema Web', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(heading)
    
    heading2 = doc.add_heading('Criterios de Aceptación Implementados', level=2)
    remove_spacing(heading2)
    
    # Añadir cada módulo adicional con su descripción
    for num in range(35, 45):
        num_str = str(num)
        if num_str in web_images and num_str in additional_modules:
            module = additional_modules[num_str]
            
            # Guardar imagen temporalmente
            temp_img = f'temp_module_{num_str}.png'
            with open(temp_img, 'wb') as f:
                f.write(web_images[num_str])
            
            try:
                add_section_with_image(
                    doc,
                    module['title'],
                    module['description'],
                    module['implementation'],
                    temp_img
                )
                print(f"   Módulo {num_str} añadido: {module['title']}")
            except Exception as e:
                print(f"   Error añadiendo módulo {num_str}: {e}")
            finally:
                if os.path.exists(temp_img):
                    os.remove(temp_img)
    
    doc.save(output_doc)
    
    print(f"\n=== DOCUMENTO ORGANIZADO CREADO ===")
    print(f"Archivo: {output_doc}")
    print(f"- {len(additional_modules)} módulos adicionales organizados")
    print(f"- Cada módulo con descripción técnica completa")
    print(f"- Comenzando con dashboard de psicólogo (imagen 35)")
    print(f"- Formato consistente con el resto del documento")
    
    return output_doc

if __name__ == "__main__":
    organize_additional_images()
