from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import os
import shutil

def create_complete_web_document():
    """Crea documento con todas las 44 imágenes del documento web"""
    
    original_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp\INFORME_FINAL_INTEGRADO.docx'
    web_images_doc = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.docx'
    output_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_COMPLETO.docx'
    
    print("1. Copiando documento original...")
    shutil.copy2(original_doc, output_doc)
    
    # Modificar texto de Android a Web
    replacements = {
        'Android': 'Web',
        'android': 'web',
        'APK': 'Web',
        'apk': 'web',
        'Android Studio': 'HTML5/CSS3/JavaScript',
        'Kotlin': 'JavaScript',
        'Java': 'JavaScript',
        'XML': 'HTML',
        'Gradle': 'NPM',
        'Room Database': 'IndexedDB',
        'SQLite': 'IndexedDB',
        'SharedPreferences': 'LocalStorage',
        'Retrofit': 'Fetch API',
        'Activity': 'Page',
        'aplicación móvil': 'aplicación web',
        'app móvil': 'app web',
        'móvil': 'web',
        'native': 'web',
        'nativo': 'web',
        'Studio': 'VS Code',
        'Emulator': 'Browser',
        'Manifest': 'Configuración',
        'build.gradle': 'package.json',
        'styles.xml': 'CSS',
        'res/': 'assets/',
        '.kt': '.js',
        '.java': '.js',
        '.xml': '.html',
    }
    
    print("2. Modificando texto...")
    doc = Document(output_doc)
    
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)
    
    # Modificar título
    for paragraph in doc.paragraphs:
        if 'Android' in paragraph.text and ('INFORME' in paragraph.text or 'Funcionamiento' in paragraph.text):
            paragraph.text = paragraph.text.replace('Android', 'Web')
            if 'Wellness Mental App' in paragraph.text:
                paragraph.text = paragraph.text.replace('Wellness Mental App', 'Wellness Mental App Web')
    
    doc.save(output_doc)
    
    print("3. Procesando imágenes...")
    # Extraer todas las imágenes del documento web
    web_images = {}
    with zipfile.ZipFile(web_images_doc, 'r') as z:
        for f in z.namelist():
            if f.startswith('word/media/image'):
                num = f.replace('word/media/image', '').replace('.png', '')
                web_images[num] = z.read(f)
    
    print(f"   Imágenes web extraídas: {len(web_images)}")
    
    # Reemplazar imágenes y añadir las faltantes
    temp = output_doc + '.tmp'
    with zipfile.ZipFile(output_doc, 'r') as z1:
        with zipfile.ZipFile(temp, 'w', zipfile.ZIP_DEFLATED) as z2:
            replaced_count = 0
            added_count = 0
            
            for info in z1.infolist():
                data = z1.read(info.filename)
                if info.filename.startswith('word/media/image'):
                    num = info.filename.replace('word/media/image', '').replace('.png', '')
                    if num in web_images:
                        z2.writestr(info.filename, web_images[num])
                        replaced_count += 1
                    else:
                        z2.writestr(info, data)
                else:
                    z2.writestr(info, data)
            
            # Añadir imágenes adicionales del documento web (image35-44)
            for num in range(35, 45):
                num_str = str(num)
                if num_str in web_images:
                    new_filename = f'word/media/image{num_str}.png'
                    z2.writestr(new_filename, web_images[num_str])
                    added_count += 1
                    print(f"   Añadida imagen adicional: image{num_str}.png")
            
            print(f"   Imágenes reemplazadas: {replaced_count}")
            print(f"   Imágenes adicionales añadidas: {added_count}")
    
    os.remove(output_doc)
    shutil.move(temp, output_doc)
    
    # Añadir las imágenes adicionales al documento usando python-docx
    print("4. Añadiendo imágenes adicionales al documento...")
    doc = Document(output_doc)
    
    # Encontrar el final del documento y añadir las imágenes adicionales
    doc.add_page_break()
    
    heading = doc.add_heading('Imágenes Adicionales de Módulos Web', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Añadir las imágenes adicionales (35-44)
    for i in range(35, 45):
        num_str = str(i)
        if num_str in web_images:
            # Guardar imagen temporalmente
            temp_img = f'temp_image_{num_str}.png'
            with open(temp_img, 'wb') as f:
                f.write(web_images[num_str])
            
            try:
                doc.add_heading(f'Imagen Adicional {num_str}', level=2)
                doc.add_picture(temp_img, width=Inches(6))  # Tamaño más grande para mejor visibilidad
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Espacio
                print(f"   Imagen {num_str} añadida al documento")
            except Exception as e:
                print(f"   Error añadiendo imagen {num_str}: {e}")
            finally:
                if os.path.exists(temp_img):
                    os.remove(temp_img)
    
    doc.save(output_doc)
    
    print(f"5. Documento final creado: {output_doc}")
    print("   - Texto Android -> Web")
    print(f"   - {len(web_images)} imágenes totales incluidas")
    print("   - Imágenes con tamaño mejorado (6 pulgadas)")
    print("   - Todas las imágenes adicionales de dashboard psicólogo incluidas")
    
    return output_doc

if __name__ == "__main__":
    create_complete_web_document()
