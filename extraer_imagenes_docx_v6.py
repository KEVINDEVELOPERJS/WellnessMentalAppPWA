from docx import Document
import os

def extract_images_from_docx_debug(docx_path, output_folder):
    """Extrae imágenes de un documento DOCX con depuración detallada"""
    os.makedirs(output_folder, exist_ok=True)
    
    try:
        doc = Document(docx_path)
        
        print(f"Document opened successfully")
        print(f"Number of paragraphs: {len(doc.paragraphs)}")
        print(f"Number of tables: {len(doc.tables)}")
        
        # Mostrar algunos párrafos de ejemplo
        print("\nFirst 5 paragraphs:")
        for i, para in enumerate(doc.paragraphs[:5]):
            print(f"{i}: {para.text[:100]}")
        
        extracted_count = 0
        
        # Buscar imágenes en el documento
        for i, paragraph in enumerate(doc.paragraphs):
            # Verificar si el párrafo contiene imágenes
            for run in paragraph.runs:
                # Buscar elementos de imagen en el XML
                pics = run._element.xpath('.//pic:pic')
                if pics:
                    print(f"\nFound {len(pics)} image(s) in paragraph {i}")
                    print(f"Paragraph text: {paragraph.text[:100]}")
                    
                    # Intentar extraer la imagen
                    try:
                        for shape in pics:
                            # Buscar la imagen incrustada
                            blips = shape.xpath('.//a:blip')
                            print(f"Found {len(blips)} blip(s)")
                            
                            for blip in blips:
                                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                print(f"Embed attribute: {embed}")
                                
                                if embed:
                                    # Obtener la imagen de las relaciones
                                    if embed in doc.part.rels:
                                        image_part = doc.part.rels[embed].target_part
                                        image_data = image_part.blob
                                        
                                        # Guardar la imagen
                                        image_name = f"image_{extracted_count + 1}.png"
                                        output_path = os.path.join(output_folder, image_name)
                                        
                                        with open(output_path, 'wb') as f:
                                            f.write(image_data)
                                        
                                        print(f"Imagen extraída: {image_name} ({len(image_data)} bytes)")
                                        extracted_count += 1
                                    else:
                                        print(f"Embed ID {embed} not found in relationships")
                    except Exception as e:
                        print(f"Error extrayendo imagen del párrafo {i}: {e}")
                        import traceback
                        traceback.print_exc()
        
        # También buscar en tablas
        print("\nSearching in tables...")
        for table_idx, table in enumerate(doc.tables):
            print(f"Table {table_idx} has {len(table.rows)} rows")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            pics = run._element.xpath('.//pic:pic')
                            if pics:
                                print(f"Found image in table {table_idx}, row {row_idx}, cell {cell_idx}")
                                try:
                                    for shape in pics:
                                        blips = shape.xpath('.//a:blip')
                                        for blip in blips:
                                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                            if embed:
                                                if embed in doc.part.rels:
                                                    image_part = doc.part.rels[embed].target_part
                                                    image_data = image_part.blob
                                                    
                                                    image_name = f"image_{extracted_count + 1}.png"
                                                    output_path = os.path.join(output_folder, image_name)
                                                    
                                                    with open(output_path, 'wb') as f:
                                                        f.write(image_data)
                                                    
                                                    print(f"Imagen extraída: {image_name} ({len(image_data)} bytes)")
                                                    extracted_count += 1
                                except Exception as e:
                                    print(f"Error extrayendo imagen de celda: {e}")
        
        print(f"\nTotal de imágenes extraídas: {extracted_count}")
        return extracted_count
        
    except Exception as e:
        print(f"Error procesando el archivo DOCX: {e}")
        import traceback
        traceback.print_exc()
        return 0

if __name__ == "__main__":
    docx_file = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp\INFORME_FINAL_INTEGRADO.docx'
    output_dir = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    print(f"Archivo DOCX: {docx_file}")
    print(f"Carpeta de salida: {output_dir}")
    print(f"Archivo DOCX existe: {os.path.exists(docx_file)}")
    
    count = extract_images_from_docx_debug(docx_file, output_dir)
    print(f"\nProceso completado. {count} imágenes extraídas en {output_dir}")
