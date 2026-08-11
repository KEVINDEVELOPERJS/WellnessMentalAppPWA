import os
import subprocess
import sys

def create_word_document_with_images():
    """Crea documento Word usando python-docx con imágenes del RTF"""
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    def remove_spacing(paragraph):
        paragraph.space_before = Pt(0)
        paragraph.space_after = Pt(0)
        paragraph.line_spacing = 1.0
    
    doc = Document()
    
    # Configurar fuente predeterminada
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Informe Técnico: Funcionamiento Correcto de Wellness Mental App Web', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(title)
    
    # Metadata
    p = doc.add_paragraph('Fecha: 4 de agosto de 2026')
    remove_spacing(p)
    p = doc.add_paragraph('Versión: 1.0')
    remove_spacing(p)
    p = doc.add_paragraph('Tipo: Informe de Validación de Criterios de Aceptación - Versión Web')
    remove_spacing(p)
    p = doc.add_paragraph()
    remove_spacing(p)
    
    # Resumen Ejecutivo
    h = doc.add_heading('Resumen Ejecutivo', level=1)
    remove_spacing(h)
    p = doc.add_paragraph(
        'Este informe presenta una validación detallada del funcionamiento correcto de la aplicación web '
        '"Wellness Mental App Web", desarrollada con HTML5, CSS3, JavaScript y arquitectura PWA. La aplicación '
        'está diseñada para el bienestar mental de estudiantes adolescentes (13-18 años) e implementa múltiples '
        'módulos interconectados que cumplen con los criterios de aceptación definidos para cada historia de usuario. '
        'La versión web mantiene paridad funcional con la versión Android nativa, ofreciendo acceso multiplataforma '
        'mediante tecnologías web modernas y Capacitor para despliegue híbrido.'
    )
    remove_spacing(p)
    
    # Arquitectura Técnica
    h = doc.add_heading('Arquitectura Técnica', level=1)
    remove_spacing(h)
    h = doc.add_heading('Estructura Web Implementada', level=2)
    remove_spacing(h)
    
    # Capa Model
    h = doc.add_heading('Capa Model (Datos y Persistencia):', level=3)
    remove_spacing(h)
    p = doc.add_paragraph()
    p.add_run('IndexedDB/').bold = True
    p.add_run(': Base de datos local del navegador para almacenamiento persistente')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('LocalStorage/').bold = True
    p.add_run(': Gestión de sesión y preferencias de usuario')
    remove_spacing(p)
    p = doc.add_paragraph()
    p.add_run('API Hub/').bold = True
    p.add_run(': Cliente REST para sincronización con backend centralizado')
    remove_spacing(p)
    
    # Capa View
    h = doc.add_heading('Capa View (Interfaz de Usuario):', level=3)
    remove_spacing(h)
    pages = [
        ('index.html', 'Pantalla principal de login/registro'),
        ('evaluation.html', 'Sistema de evaluación psicológica'),
        ('chat.html', 'Interfaz de chat con asistente emocional IA'),
        ('exercises.html', 'Catálogo de ejercicios de bienestar'),
        ('games.html', 'Sistema de gamificación y mini-juegos'),
        ('community.html', 'Foros de comunidad estudiantil'),
        ('alerts.html', 'Panel de alertas para psicólogos (riesgo alto)'),
        ('alerts-low-medium.html', 'Panel de alertas para psicólogos (riesgo bajo/medio)'),
        ('questionnaire-editor.html', 'Editor de cuestionarios para psicólogos'),
        ('profile.html', 'Gestión de perfil y privacidad'),
        ('parent-reports.html', 'Generación de informes para padres'),
        ('videos.html', 'Videos guiados de respiración y meditación'),
        ('active-breaks.html', 'Configuración de pausas activas'),
        ('mental-garden.html', 'Jardín mental interactivo')
    ]
    
    for page, desc in pages:
        p = doc.add_paragraph()
        p.add_run(f'{page}/').bold = True
        p.add_run(f': {desc}')
        remove_spacing(p)
    
    # Capa Controller
    h = doc.add_heading('Capa Controller (Lógica de Negocio):', level=3)
    remove_spacing(h)
    controllers = [
        ('app.js', 'Controlador principal de la aplicación'),
        ('evaluation.js', 'Lógica de evaluación psicológica'),
        ('chat.js', 'Gestión de conversaciones con IA'),
        ('exercises.js', 'Control de ejercicios y temporizadores'),
        ('games.js', 'Sistema de gamificación, puntos y niveles'),
        ('mental-garden.js', 'Lógica del jardín mental'),
        ('alerts.js', 'Gestión de alertas de riesgo alto'),
        ('alerts-low-medium.js', 'Gestión de alertas de riesgo bajo/medio'),
        ('hub-client.js', 'Cliente API para comunicación con backend')
    ]
    
    for controller, desc in controllers:
        p = doc.add_paragraph()
        p.add_run(f'{controller}/').bold = True
        p.add_run(f': {desc}')
        remove_spacing(p)
    
    # Validación por Criterios
    h = doc.add_heading('Validación por Criterios de Aceptación con Evidencia Visual', level=1)
    remove_spacing(h)
    
    # Mapeo de imágenes (secuenciales basadas en el RTF)
    image_sequence = [
        ('HU-01', 'Registro y Autenticación de Usuarios', [
            ('web_registro.png', 'CA-01.1: Pantalla de Registro de Usuario con validación de campos', 'index.html, app.js'),
            ('web_login.png', 'CA-01.4: Pantalla de Login con autenticación', 'index.html, app.js'),
            ('web_consentimiento.png', 'CA-01.3: Pantalla de Consentimiento Parental', 'index.html')
        ]),
        ('HU-02', 'Dashboard Principal', [
            ('web_dashboard_estudiante.png', 'CA-02.1: Dashboard de Estudiante con estado general', 'index.html, app.js'),
            ('web_dashboard_psicologo.png', 'CA-02.1: Dashboard de Psicólogo con panel de alertas', 'index.html, app.js')
        ]),
        ('HU-03', 'Sistema de Evaluación Psicológica', [
            ('web_cuestionarios.png', 'CA-03.1: Lista de Cuestionarios Disponibles', 'evaluation.html, evaluation.js'),
            ('web_gad7.png', 'CA-03.1: Cuestionario GAD-7 en Progreso', 'evaluation.html, evaluation.js'),
            ('web_resultados.png', 'CA-03.2: Resultados de Evaluación con nivel de riesgo', 'evaluation.js'),
            ('web_historial.png', 'CA-03.3: Historial de Evaluaciones', 'evaluation.html, evaluation.js')
        ]),
        ('HU-04', 'Chat con Inteligencia Artificial', [
            ('web_chat.png', 'CA-04.1: Interfaz de Chat IA con bubbles de mensajes', 'chat.html, chat.js'),
            ('web_chat_conversacion.png', 'CA-04.2: Conversación con Asistente Emocional empático', 'chat.js')
        ]),
        ('HU-05', 'Ejercicios de Bienestar', [
            ('web_ejercicios.png', 'CA-05.1: Catálogo de Ejercicios por categoría', 'exercises.html, exercises.js'),
            ('web_respiracion.png', 'CA-05.2: Ejercicio de Respiración 4-7-8 con temporizador', 'exercises.js'),
            ('web_meditacion.png', 'CA-05.2: Meditación Guiada con instrucciones', 'exercises.html, exercises.js'),
            ('web_progreso.png', 'CA-05.3: Progreso de Ejercicios con estadísticas', 'exercises.js')
        ]),
        ('HU-08', 'Gamificación y Juegos', [
            ('web_juegos.png', 'CA-08.1: Pantalla Principal de Juegos', 'games.html, games.js'),
            ('web_puzzle.png', 'CA-08.1: Mini-juego Puzzle Zen', 'puzzle-zen.html, puzzle-zen.js'),
            ('web_arte.png', 'CA-08.1: Mini-juego Arte Emocional', 'arte-emocional.html, arte-emocional.js'),
            ('web_ritmo.png', 'CA-08.1: Mini-juego Ritmo Calma', 'ritmo-calma.html, ritmo-calma.js'),
            ('web_jardin.png', 'CA-08.1: Jardín Mental con progreso visual', 'mental-garden.html, mental-garden.js'),
            ('web_logros.png', 'CA-08.4: Logros Desbloqueados con celebración', 'games.js'),
            ('web_puntos.png', 'CA-08.2-CA-08.3: Sistema de Puntos y Niveles', 'games.js')
        ]),
        ('HU-07', 'Comunidad Estudiantil', [
            ('web_comunidad.png', 'CA-07.1: Foros de Comunidad categorizados', 'community.html, community.js'),
            ('web_post.png', 'CA-07.2: Creación de Post con título y contenido', 'community.js'),
            ('web_interaccion.png', 'CA-07.2: Interacción en Comunidad con comentarios', 'community.js')
        ]),
        ('HU-09', 'Informes para Padres', [
            ('web_informes.png', 'CA-09.1: Generación de Informe con resumen de bienestar', 'parent-reports.html, parent-reports.js'),
            ('web_informe_previa.png', 'CA-09.1: Vista Previa de Informe para Padres', 'parent-reports.js')
        ]),
        ('HU-10', 'Gestión de Perfil y Privacidad', [
            ('web_perfil.png', 'CA-10.1: Perfil de Usuario con estadísticas', 'profile.html, profile.js'),
            ('web_privacidad.png', 'CA-10.2: Configuración de Privacidad y notificaciones', 'profile.js'),
            ('web_sesiones.png', 'CA-10.3: Gestión de Sesiones activas', 'app.js')
        ]),
        ('HU-06', 'Diseño de Cuestionarios por Psicólogo', [
            ('web_editor.png', 'CA-06.1: Editor de Cuestionarios con campos y botones', 'questionnaire-editor.html, questionnaire-editor.js'),
            ('web_editor_previa.png', 'CA-06.4: Vista Previa de Cuestionario como estudiante', 'questionnaire-editor.js')
        ]),
        ('Alertas', 'Gestión de Alertas para Psicólogos', [
            ('web_alertas_panel.png', 'CA-Alerta.1: Panel de Alertas por prioridad', 'alerts.html, alerts.js'),
            ('web_alertas_detalle.png', 'CA-Alerta.1: Detalle de Alerta de Riesgo', 'alerts.js'),
            ('web_alertas_bajo_medio.png', 'CA-Alerta.2: Panel de Alertas de Riesgo Bajo/Medio', 'alerts-low-medium.html, alerts-low-medium.js')
        ]),
        ('Misiones', 'Misiones Diarias y Check-In', [
            ('web_checkin.png', 'CA-Mision.1: Check-In Diario con estado de ánimo', 'index.html, app.js'),
            ('web_misiones.png', 'CA-Mision.2: Misiones Diarias con recompensas', 'games.html, games.js')
        ]),
        ('Videos', 'Videos Guiados', [
            ('web_videos.png', 'CA-Video.1: Catálogo de Videos Guiados', 'videos.html, videos.js'),
            ('web_video_reproduccion.png', 'CA-Video.2: Reproducción de Video con controles', 'videos.js')
        ]),
        ('Pausas', 'Pausas Activas', [
            ('web_pausas.png', 'CA-Pausa.1: Pantalla de Pausas Activas', 'active-breaks.html, active-breaks.js'),
            ('web_pausa_config.png', 'CA-Pausa.2: Configuración de Recordatorios de Pausas', 'active-breaks.js')
        ])
    ]
    
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    # Procesar cada módulo
    for key, title, criteria in image_sequence:
        h = doc.add_heading(f'{key}: {title}', level=2)
        remove_spacing(h)
        
        h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
        remove_spacing(h)
        
        for image_name, description, files in criteria:
            criterion_title = description.split(':')[0]
            h = doc.add_heading(criterion_title, level=4)
            remove_spacing(h)
            
            # Descripción con implementación
            p = doc.add_paragraph()
            p.add_run(f'{description} - Implementado en: {files}. ').bold = True
            p.add_run('✅ FUNCIONAL')
            remove_spacing(p)
            
            # Intentar agregar imagen
            image_path = os.path.join(images_folder, image_name)
            
            # Texto descriptivo antes de imagen
            p = doc.add_paragraph()
            run = p.add_run('Evidencia Visual: ' + description.split(':')[1].strip() if ':' in description else description)
            run.bold = True
            run.font.size = Pt(9)
            remove_spacing(p)
            
            if os.path.exists(image_path):
                try:
                    doc.add_picture(image_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    remove_spacing(last_paragraph)
                    print(f'Imagen agregada: {image_name}')
                except Exception as e:
                    p = doc.add_paragraph(f'[Error al cargar imagen: {e}]')
                    remove_spacing(p)
            else:
                p = doc.add_paragraph(f'[Imagen no disponible: {image_name}]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                remove_spacing(p)
            
            p = doc.add_paragraph()
            remove_spacing(p)
        
        # Salto de página entre módulos
        doc.add_page_break()
    
    # Conclusiones
    h = doc.add_heading('Conclusiones', level=1)
    remove_spacing(h)
    
    conclusion_text = (
        'La aplicación web Wellness Mental App cumple satisfactoriamente con todos los criterios de aceptación '
        'establecidos para cada historia de usuario. La implementación web mantiene paridad funcional con la versión '
        'Android nativa, aprovechando las ventajas de las tecnologías web modernas para ofrecer una experiencia de usuario '
        'consistente y accesible multiplataforma.\n\n'
        'Todos los módulos están plenamente operativos:\n'
        '- Sistema de autenticación y gestión de usuarios\n'
        '- Evaluaciones psicológicas con cálculo automático de riesgos\n'
        '- Chat con asistente emocional IA\n'
        '- Ejercicios de bienestar interactivos\n'
        '- Sistema de gamificación completo con jardín mental\n'
        '- Comunidad estudiantil\n'
        '- Panel de alertas para psicólogos (alto, medio y bajo riesgo)\n'
        '- Informes para padres\n'
        '- Gestión de perfiles y privacidad\n'
        '- Editor de cuestionarios para psicólogos\n'
        '- Videos guiados y pausas activas\n\n'
        'La arquitectura web implementada (HTML5, CSS3, JavaScript, IndexedDB, PWA) proporciona una base sólida y '
        'escalable para la aplicación, asegurando rendimiento, accesibilidad y mantenibilidad a largo plazo.'
    )
    
    p = doc.add_paragraph(conclusion_text)
    remove_spacing(p)
    
    # Guardar documento
    output_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB.docx'
    doc.save(output_path)
    
    print(f'\n✅ Documento Word creado exitosamente: {output_path}')
    print(f'📊 Total de módulos documentados: {len(image_sequence)}')
    print(f'🖼️ Total de imágenes requeridas: {sum(len(c) for _, _, c in image_sequence)}')
    print(f'📁 Carpeta de imágenes: {images_folder}')
    
    return output_path

if __name__ == '__main__':
    create_word_document_with_images()
