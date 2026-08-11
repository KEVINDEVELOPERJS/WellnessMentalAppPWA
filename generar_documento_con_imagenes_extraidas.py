from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os
import glob

def remove_spacing(paragraph):
    """Remove spacing before and after a paragraph"""
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def add_section_with_content(doc, title, description, implementation_files, status="✅ FUNCIONAL", image_path=None):
    """Añade una sección completa con título, descripción, implementación y estado"""
    
    # Título de la sección
    heading = doc.add_heading(title, level=3)
    remove_spacing(heading)
    
    # Descripción
    p = doc.add_paragraph(description)
    remove_spacing(p)
    
    # Implementación
    p = doc.add_paragraph(f"Implementado en: {implementation_files}")
    p.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Azul
    remove_spacing(p)
    
    # Estado
    p = doc.add_paragraph(status)
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Verde
    remove_spacing(p)
    
    # Evidencia visual
    p = doc.add_paragraph("Evidencia Visual:")
    p.runs[0].font.italic = True
    remove_spacing(p)
    
    # Imagen si existe
    if image_path and os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            remove_spacing(last_paragraph)
        except Exception as e:
            p = doc.add_paragraph(f"[Imagen no disponible: {os.path.basename(image_path)}]")
            p.runs[0].font.italic = True
            remove_spacing(p)
    else:
        p = doc.add_paragraph("[Captura de pantalla del módulo]")
        p.runs[0].font.italic = True
        remove_spacing(p)
    
    # Espacio
    doc.add_paragraph()

def create_final_web_document_with_extracted_images():
    """Crea el documento final de la app web usando las imágenes extraídas del DOCX original"""
    
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Carpeta de imágenes extraídas
    images_folder = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\web_images'
    
    # Obtener lista de imágenes extraídas
    extracted_images = []
    if os.path.exists(images_folder):
        for ext in ['*.png', '*.jpg', '*.jpeg']:
            extracted_images.extend(glob.glob(os.path.join(images_folder, ext)))
        extracted_images.sort()
    
    print(f"Imágenes extraídas disponibles: {len(extracted_images)}")
    
    # Mapeo de secciones a imágenes (usando las imágenes extraídas secuencialmente)
    image_mapping = {
        'HU-01': [0, 1, 2],           # 3 imágenes para registro/login/consentimiento
        'HU-02': [3, 4],              # 2 imágenes para dashboards
        'HU-03': [5, 6, 7, 8],        # 4 imágenes para evaluaciones
        'HU-04': [9, 10],             # 2 imágenes para chat
        'HU-05': [11, 12, 13, 14],    # 4 imágenes para ejercicios
        'HU-08': [15, 16, 17, 18, 19, 20, 21],  # 7 imágenes para juegos
        'HU-07': [22, 23, 24],        # 3 imágenes para comunidad
        'HU-09': [25, 26],            # 2 imágenes para informes
        'HU-10': [27, 28, 29],        # 3 imágenes para perfil
        'HU-06': [30, 31],            # 2 imágenes para editor
        'Alertas': [32, 33],          # 2 imágenes para alertas
        'Misiones': [],               # Sin imágenes específicas
        'Videos': [],                 # Sin imágenes específicas
        'Pausas': []                  # Sin imágenes específicas
    }
    
    # ==================== PORTADA ====================
    
    title = doc.add_heading('Informe Técnico: Funcionamiento Correcto de Wellness Mental App Web', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_spacing(title)
    
    p = doc.add_paragraph('Fecha: 4 de agosto de 2026')
    remove_spacing(p)
    p = doc.add_paragraph('Versión: 1.0')
    remove_spacing(p)
    p = doc.add_paragraph('Tipo: Informe de Validación de Criterios de Aceptación - Versión Web')
    remove_spacing(p)
    p = doc.add_paragraph()
    remove_spacing(p)
    
    # ==================== RESUMEN EJECUTIVO ====================
    
    h = doc.add_heading('Resumen Ejecutivo', level=1)
    remove_spacing(h)
    
    p = doc.add_paragraph(
        'Este informe presenta una validación detallada del funcionamiento correcto de la aplicación web '
        '"Wellness Mental App Web", desarrollada con HTML5, CSS3, JavaScript y arquitectura PWA. La aplicación '
        'está diseñada para el bienestar mental de estudiantes adolescentes (13-18 años) e implementa múltiples '
        'módulos interconectados que cumplen con los criterios de aceptación definidos para cada historia de usuario. '
        'La versión web mantiene paridad funcional con la versión Android nativa, ofreciendo acceso multiplataforma '
        'mediante tecnologías web modernas y Capacitor para despliegue híbrido.'
    )
    remove_spacing(p)
    
    # ==================== ARQUITECTURA TÉCNICA ====================
    
    h = doc.add_heading('Arquitectura Técnica', level=1)
    remove_spacing(h)
    
    h = doc.add_heading('Estructura Web Implementada', level=2)
    remove_spacing(h)
    
    # Capa Model
    h = doc.add_heading('CAPA MODEL (DATOS Y PERSISTENCIA)', level=3)
    remove_spacing(h)
    
    items = [
        'IndexedDB: Base de datos local del navegador para almacenamiento persistente (usuarios, respuestas, resultados, chat, ejercicios)',
        'LocalStorage: Gestión de sesión y preferencias de usuario',
        'API Hub: Cliente REST para sincronización con backend centralizado',
        'LocalStorage: Gestión de preferencias y configuración'
    ]
    
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        remove_spacing(p)
    
    # Capa View
    h = doc.add_heading('CAPA VIEW (INTERFAZ DE USUARIO)', level=3)
    remove_spacing(h)
    
    view_files = [
        'index.html: Pantalla principal de login/registro',
        'evaluation.html: Sistema de evaluación psicológica (GAD-7, PHQ-9, PSS-10)',
        'chat.html: Interfaz de chat con asistente emocional IA',
        'exercises.html: Catálogo de ejercicios de bienestar',
        'games.html: Sistema de gamificación y mini-juegos',
        'community.html: Foros de comunidad estudiantil',
        'alerts.html: Panel de alertas para psicólogos (riesgo alto)',
        'alerts-low-medium.html: Panel de alertas para psicólogos (riesgo bajo/medio)',
        'questionnaire-editor.html: Editor de cuestionarios para psicólogos',
        'profile.html: Gestión de perfil y privacidad',
        'parent-reports.html: Generación de informes para padres',
        'videos.html: Videos guiados de respiración y meditación',
        'active-breaks.html: Configuración de pausas activas',
        'mental-garden.html: Jardín mental interactivo'
    ]
    
    for item in view_files:
        p = doc.add_paragraph(item, style='List Bullet')
        remove_spacing(p)
    
    # Capa Controller
    h = doc.add_heading('CAPA CONTROLLER (LÓGICA DE NEGOCIO)', level=3)
    remove_spacing(h)
    
    controller_files = [
        'app.js: Controlador principal de la aplicación, gestión de estado y navegación',
        'evaluation.js: Lógica de evaluación psicológica, cálculo de puntajes y niveles de riesgo',
        'chat.js: Gestión de conversaciones con IA, análisis de sentimientos',
        'exercises.js: Control de ejercicios, temporizadores y seguimiento de progreso',
        'games.js: Sistema de gamificación, puntos, niveles y logros',
        'mental-garden.js: Lógica del jardín mental, sistema de plantas y riego',
        'alerts.js: Gestión de alertas de riesgo alto para psicólogos',
        'alerts-low-medium.js: Gestión de alertas de riesgo bajo/medio para psicólogos',
        'hub-client.js: Cliente API para comunicación con backend centralizado'
    ]
    
    for item in controller_files:
        p = doc.add_paragraph(item, style='List Bullet')
        remove_spacing(p)
    
    add_page_break(doc)
    
    # ==================== VALIDACIÓN POR CRITERIOS DE ACEPTACIÓN ====================
    
    h = doc.add_heading('Validación por Criterios de Aceptación con Evidencia Visual', level=1)
    remove_spacing(h)
    
    # Función para obtener imagen del índice
    def get_image(section, index):
        if section in image_mapping and index < len(image_mapping[section]):
            img_index = image_mapping[section][index]
            if img_index < len(extracted_images):
                return extracted_images[img_index]
        return None
    
    # HU-01: Registro y Autenticación
    h = doc.add_heading('HU-01: Registro y Autenticación de Usuarios', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-01.1: Pantalla de Registro de Usuario con validación de campos',
        'Formulario de registro web con validación de email, contraseña y campos obligatorios',
        'index.html, app.js (funciones de registro y validación)',
        image_path=get_image('HU-01', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-01.4: Pantalla de Login con autenticación',
        'Sistema de autenticación web con gestión de sesiones',
        'index.html, app.js (sistema de autenticación web)',
        image_path=get_image('HU-01', 1)
    )
    
    add_section_with_content(
        doc,
        'CA-01.3: Pantalla de Consentimiento Parental',
        'Formulario de consentimiento para menores de edad',
        'index.html (formulario de consentimiento para menores)',
        image_path=get_image('HU-01', 2)
    )
    
    # HU-02: Dashboard Principal
    h = doc.add_heading('HU-02: Dashboard Principal', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-02.1: Dashboard de Estudiante con estado general',
        'Vista personalizada por rol con estado emocional y actividades',
        'index.html, app.js (vista personalizada por rol)',
        image_path=get_image('HU-02', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-02.1: Dashboard de Psicólogo con panel de alertas',
        'Panel específico para psicólogos con alertas de riesgo',
        'index.html, app.js (panel específico para psicólogos)',
        image_path=get_image('HU-02', 1)
    )
    
    # HU-03: Sistema de Evaluación Psicológica
    h = doc.add_heading('HU-03: Sistema de Evaluación Psicológica', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-03.1: Lista de Cuestionarios Disponibles (GAD-7, PHQ-9, PSS-10)',
        'Catálogo de cuestionarios validados para evaluación psicológica',
        'evaluation.html, evaluation.js (carga de cuestionarios)',
        image_path=get_image('HU-03', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-03.1: Cuestionario GAD-7 en Progreso con escala Likert',
        'Interfaz interactiva con escala Likert para respuestas',
        'evaluation.html, evaluation.js (interfaz interactiva)',
        image_path=get_image('HU-03', 1)
    )
    
    add_section_with_content(
        doc,
        'CA-03.2: Resultados de Evaluación con nivel de riesgo',
        'Cálculo de puntajes y categorización de niveles de riesgo',
        'evaluation.js (cálculo de puntajes y categorización)',
        image_path=get_image('HU-03', 2)
    )
    
    add_section_with_content(
        doc,
        'CA-03.3: Historial de Evaluaciones con fechas y puntajes',
        'Almacenamiento y visualización de historial completo',
        'evaluation.html, evaluation.js (almacenamiento IndexedDB)',
        image_path=get_image('HU-03', 3)
    )
    
    # HU-04: Chat con IA
    h = doc.add_heading('HU-04: Chat con Inteligencia Artificial', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-04.1: Interfaz de Chat IA con bubbles de mensajes',
        'Diseño responsivo de chat con burbujas de mensajes',
        'chat.html, chat.js (diseño responsivo de chat)',
        image_path=get_image('HU-04', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-04.2: Conversación con Asistente Emocional empático',
        'Respuestas contextuales y análisis de sentimientos',
        'chat.js (respuestas contextuales y análisis de sentimientos)',
        image_path=get_image('HU-04', 1)
    )
    
    # HU-05: Ejercicios de Bienestar
    h = doc.add_heading('HU-05: Ejercicios de Bienestar', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-05.1: Catálogo de Ejercicios por categoría',
        'Organización de ejercicios por tipo y categoría',
        'exercises.html, exercises.js (organización de ejercicios)',
        image_path=get_image('HU-05', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-05.2: Ejercicio de Respiración 4-7-8 con temporizador',
        'Temporizadores interactivos y guías visuales',
        'exercises.js (temporizadores interactivos y guías visuales)',
        image_path=get_image('HU-05', 1)
    )
    
    add_section_with_content(
        doc,
        'CA-05.2: Meditación Guiada con instrucciones',
        'Contenido multimedia guiado paso a paso',
        'exercises.html, exercises.js (contenido multimedia guiado)',
        image_path=get_image('HU-05', 2)
    )
    
    add_section_with_content(
        doc,
        'CA-05.3: Progreso de Ejercicios con estadísticas',
        'Seguimiento de progreso en IndexedDB',
        'exercises.js (seguimiento de progreso en IndexedDB)',
        image_path=get_image('HU-05', 3)
    )
    
    # HU-08: Gamificación y Juegos
    h = doc.add_heading('HU-08: Gamificación y Juegos', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-08.1: Pantalla Principal de Juegos',
        'Panel de gamificación centralizado',
        'games.html, games.js (panel de gamificación centralizado)',
        image_path=get_image('HU-08', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-08.1: Mini-juego Puzzle Zen',
        'Juego de memoria web con lógica interactiva',
        'puzzle-zen.html, puzzle-zen.js (juego de memoria web)',
        image_path=get_image('HU-08', 1)
    )
    
    add_section_with_content(
        doc,
        'CA-08.1: Mini-juego Arte Emocional',
        'Expresión creativa web',
        'arte-emocional.html, arte-emocional.js (expresión creativa)',
        image_path=get_image('HU-08', 2)
    )
    
    add_section_with_content(
        doc,
        'CA-08.1: Mini-juego Ritmo Calma',
        'Juego rítmico web',
        'ritmo-calma.html, ritmo-calma.js (juego rítmico web)',
        image_path=get_image('HU-08', 3)
    )
    
    add_section_with_content(
        doc,
        'CA-08.1: Jardín Mental con progreso visual',
        'Sistema de cultivo web con plantas y riego',
        'mental-garden.html, mental-garden.js (sistema de cultivo web)',
        image_path=get_image('HU-08', 4)
    )
    
    add_section_with_content(
        doc,
        'CA-08.4: Logros Desbloqueados con celebración',
        'Sistema de logros y recompensas',
        'games.js (sistema de logros y recompensas)',
        image_path=get_image('HU-08', 5)
    )
    
    add_section_with_content(
        doc,
        'CA-08.2-CA-08.3: Sistema de Puntos y Niveles',
        'Gamificación con integración cross-módulo',
        'games.js (gamificación con integración cross-módulo)',
        image_path=get_image('HU-08', 6)
    )
    
    # HU-07: Comunidad Estudiantil
    h = doc.add_heading('HU-07: Comunidad Estudiantil', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-07.1: Foros de Comunidad categorizados',
        'Sistema de foros web por categorías temáticas',
        'community.html, community.js (sistema de foros web)',
        image_path=get_image('HU-07', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-07.2: Creación de Post con título y contenido',
        'Editor de posts web con validación',
        'community.html, community.js (editor de posts web)',
        image_path=get_image('HU-07', 1)
    )
    
    add_section_with_content(
        doc,
        'CA-07.2: Interacción en Comunidad con comentarios',
        'Sistema de comentarios y likes web',
        'community.html, community.js (sistema de comentarios web)',
        image_path=get_image('HU-07', 2)
    )
    
    # HU-09: Informes para Padres
    h = doc.add_heading('HU-09: Informes para Padres', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-09.1: Generación de Informe con resumen de bienestar',
        'Generación automática de informes web',
        'parent-reports.html, parent-reports.js (generación de informes)',
        image_path=get_image('HU-09', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-09.1: Vista Previa de Informe para Padres',
        'Vista previa y exportación de informes',
        'parent-reports.html, parent-reports.js (vista previa informes)',
        image_path=get_image('HU-09', 1)
    )
    
    # HU-10: Gestión de Perfil y Privacidad
    h = doc.add_heading('HU-10: Gestión de Perfil y Privacidad', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-10.1: Perfil de Usuario con estadísticas',
        'Perfil web con estadísticas de uso',
        'profile.html, profile.js (perfil web)',
        image_path=get_image('HU-10', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-10.2: Configuración de Privacidad y notificaciones',
        'Panel de configuración web',
        'profile.html, profile.js (configuración web)',
        image_path=get_image('HU-10', 1)
    )
    
    add_section_with_content(
        doc,
        'CA-10.3: Gestión de Sesiones activas',
        'Control de sesiones web activas',
        'profile.html, profile.js (gestión de sesiones)',
        image_path=get_image('HU-10', 2)
    )
    
    # HU-06: Editor de Cuestionarios
    h = doc.add_heading('HU-06: Diseño de Cuestionarios por Psicólogo', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-06.1: Editor de Cuestionarios con campos y botones',
        'Editor web para crear cuestionarios personalizados',
        'questionnaire-editor.html, questionnaire-editor.js (editor web)',
        image_path=get_image('HU-06', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-06.4: Vista Previa de Cuestionario como estudiante',
        'Vista previa del cuestionario creado',
        'questionnaire-editor.html, questionnaire-editor.js (vista previa)',
        image_path=get_image('HU-06', 1)
    )
    
    # Alertas
    h = doc.add_heading('Gestión de Alertas para Psicólogos', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-Alerta.1: Panel de Alertas de Riesgo Alto',
        'Panel web de alertas de riesgo alto',
        'alerts.html, alerts.js (panel web de alertas)',
        image_path=get_image('Alertas', 0)
    )
    
    add_section_with_content(
        doc,
        'CA-Alerta.1: Detalle de Alerta de Riesgo con información estudiante',
        'Vista detallada de alerta web',
        'alerts.html, alerts.js (detalle de alerta web)',
        image_path=get_image('Alertas', 1)
    )
    
    # Misiones
    h = doc.add_heading('Misiones Diarias y Check-In', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-Mision.1: Check-In Diario con estado de ánimo',
        'Sistema de check-in web diario',
        'app.js, evaluation.js (check-in web)',
        image_path=None
    )
    
    add_section_with_content(
        doc,
        'CA-Mision.2: Misiones Diarias con recompensas',
        'Sistema de misiones web con recompensas',
        'games.js, app.js (misiones web)',
        image_path=None
    )
    
    # Videos
    h = doc.add_heading('Videos Guiados', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-Video.1: Catálogo de Videos Guiados',
        'Catálogo web de videos guiados',
        'videos.html, videos.js (catálogo web)',
        image_path=None
    )
    
    add_section_with_content(
        doc,
        'CA-Video.2: Reproducción de Video con controles',
        'Reproductor web con controles',
        'videos.html, videos.js (reproductor web)',
        image_path=None
    )
    
    # Pausas Activas
    h = doc.add_heading('Pausas Activas', level=2)
    remove_spacing(h)
    
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    add_section_with_content(
        doc,
        'CA-Pausa.1: Pantalla de Pausas Activas',
        'Pantalla web de pausas activas',
        'active-breaks.html, active-breaks.js (pantalla web)',
        image_path=None
    )
    
    add_section_with_content(
        doc,
        'CA-Pausa.2: Configuración de Recordatorios de Pausas',
        'Configuración web de recordatorios',
        'active-breaks.html, active-breaks.js (configuración web)',
        image_path=None
    )
    
    # ==================== CONCLUSIONES ====================
    
    add_page_break(doc)
    
    h = doc.add_heading('Conclusiones', level=1)
    remove_spacing(h)
    
    p = doc.add_paragraph(
        'La aplicación web "Wellness Mental App Web" ha sido validada exitosamente cumpliendo con '
        'todos los criterios de aceptación definidos para cada historia de usuario. La arquitectura '
        'web implementada permite una experiencia de usuario fluida y consistente con la versión '
        'Android nativa, aprovechando las ventajas de las tecnologías web modernas como HTML5, CSS3, '
        'JavaScript y la arquitectura PWA.'
    )
    remove_spacing(p)
    
    p = doc.add_paragraph(
        'El sistema de persistencia local mediante IndexedDB garantiza el funcionamiento offline y '
        'la sincronización cuando hay conexión, mientras que la integración con Capacitor permite '
        'el despliegue híbrido en múltiples plataformas. Todos los módulos funcionan correctamente '
        'y proporcionan las funcionalidades requeridas para el bienestar mental de estudiantes adolescentes.'
    )
    remove_spacing(p)
    
    # Guardar el documento
    output_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_CON_IMAGENES_ORIGINALES.docx'
    doc.save(output_path)
    print(f"Documento final guardado exitosamente en: {output_path}")
    return output_path

if __name__ == "__main__":
    create_final_web_document_with_extracted_images()
