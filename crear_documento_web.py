from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

def remove_spacing(paragraph):
    """Remove spacing before and after a paragraph"""
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def create_web_document():
    """Crea el documento Word basado en el contenido web existente"""
    
    # Leer el contenido del informe web
    content_file = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_WEB_COMPLETO.txt'
    
    with open(content_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Procesar el contenido línea por línea
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # Títulos principales (con ========)
        if line.startswith('=') and line.endswith('='):
            title_text = line.replace('=', '').strip()
            heading = doc.add_heading(title_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            remove_spacing(heading)
        
        # Subtítulos (con --------)
        elif line.startswith('-') and line.endswith('-'):
            subtitle_text = line.replace('-', '').strip()
            heading = doc.add_heading(subtitle_text, level=2)
            remove_spacing(heading)
        
        # Líneas que parecen títulos de módulos (HU-XX: ...)
        elif line.startswith('HU-') or line.startswith('CA-'):
            if line.startswith('HU-'):
                heading = doc.add_heading(line, level=2)
            else:
                heading = doc.add_heading(line, level=3)
            remove_spacing(heading)
        
        # Líneas con indicadores de funcionalidad (✅)
        elif '✅' in line or 'FUNCIONAL' in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Verde
            remove_spacing(p)
        
        # Líneas de evidencia visual
        elif 'Evidencia Visual:' in line or 'Captura:' in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.italic = True
            remove_spacing(p)
        
        # Líneas de implementación
        elif line.startswith('Implementado en:'):
            p = doc.add_paragraph(line)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # Azul
            remove_spacing(p)
        
        # Texto normal
        else:
            p = doc.add_paragraph(line)
            remove_spacing(p)
    
    # Guardar el documento
    output_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB.docx'
    doc.save(output_path)
    print(f"Documento guardado exitosamente en: {output_path}")
    return output_path

if __name__ == "__main__":
    create_web_document()
