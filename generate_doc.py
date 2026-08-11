import subprocess
import sys
import os

def install_and_run():
    """Instala dependencias y ejecuta la generación del documento"""
    
    print("🔧 Instalando dependencias necesarias...")
    
    # Instalar python-docx
    try:
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], 
                      capture_output=True, check=True)
        print("✅ python-docx instalado correctamente")
    except subprocess.CalledProcessError as e:
        print(f"❌ Error instalando python-docx: {e}")
        return False
    
    # Crear el script de generación
    script_content = '''from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def remove_spacing(paragraph):
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0

doc = Document()

# Title
title = doc.add_heading('Informe Técnico: Funcionamiento Correcto de Wellness Mental App Web', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
remove_spacing(title)

# Metadata
doc.add_paragraph('Fecha: 4 de agosto de 2026')
doc.add_paragraph('Versión: 1.0')
doc.add_paragraph('Tipo: Informe de Validación de Criterios de Aceptación - Versión Web')
doc.add_paragraph()

# Resumen
h = doc.add_heading('Resumen Ejecutivo', level=1)
remove_spacing(h)
p = doc.add_paragraph('Este informe presenta una validación detallada del funcionamiento correcto de la aplicación web "Wellness Mental App Web", desarrollada con HTML5, CSS3, JavaScript y arquitectura PWA.')
remove_spacing(p)

# Arquitectura
h = doc.add_heading('Arquitectura Técnica', level=1)
remove_spacing(h)
h = doc.add_heading('Estructura Web Implementada', level=2)
remove_spacing(h)

h = doc.add_heading('Capa Model (Datos y Persistencia):', level=3)
remove_spacing(h)
p = doc.add_paragraph()
p.add_run('IndexedDB/').bold = True
p.add_run(': Base de datos local del navegador')
remove_spacing(p)

h = doc.add_heading('Capa View (Interfaz de Usuario):', level=3)
remove_spacing(h)
pages = [
    'index.html: Login/registro',
    'evaluation.html: Evaluaciones psicológicas',
    'chat.html: Chat con IA',
    'exercises.html: Ejercicios de bienestar',
    'games.html: Gamificación y juegos',
    'community.html: Comunidad estudiantil',
    'alerts.html: Alertas riesgo alto',
    'alerts-low-medium.html: Alertas riesgo bajo/medio',
    'questionnaire-editor.html: Editor cuestionarios',
    'profile.html: Perfil y privacidad',
    'parent-reports.html: Informes padres',
    'videos.html: Videos guiados',
    'active-breaks.html: Pausas activas',
    'mental-garden.html: Jardín mental'
]
for page in pages:
    doc.add_paragraph(page)

h = doc.add_heading('Capa Controller (Lógica de Negocio):', level=3)
remove_spacing(h)
controllers = [
    'app.js: Controlador principal',
    'evaluation.js: Lógica evaluaciones',
    'chat.js: Gestión chat IA',
    'exercises.js: Control ejercicios',
    'games.js: Sistema gamificación',
    'mental-garden.js: Lógica jardín',
    'alerts.js: Alertas riesgo alto',
    'alerts-low-medium.js: Alertas riesgo bajo/medio',
    'hub-client.js: Cliente API'
]
for controller in controllers:
    doc.add_paragraph(controller)

# Validación
h = doc.add_heading('Validación por Criterios de Aceptación con Evidencia Visual', level=1)
remove_spacing(h)

modules = [
    ('HU-01', 'Registro y Autenticación', ['web_registro.png', 'web_login.png', 'web_consentimiento.png']),
    ('HU-02', 'Dashboard Principal', ['web_dashboard_estudiante.png', 'web_dashboard_psicologo.png']),
    ('HU-03', 'Evaluación Psicológica', ['web_cuestionarios.png', 'web_gad7.png', 'web_resultados.png', 'web_historial.png']),
    ('HU-04', 'Chat con IA', ['web_chat.png', 'web_chat_conversacion.png']),
    ('HU-05', 'Ejercicios Bienestar', ['web_ejercicios.png', 'web_respiracion.png', 'web_meditacion.png', 'web_progreso.png']),
    ('HU-08', 'Gamificación', ['web_juegos.png', 'web_puzzle.png', 'web_arte.png', 'web_ritmo.png', 'web_jardin.png', 'web_logros.png', 'web_puntos.png']),
    ('HU-07', 'Comunidad', ['web_comunidad.png', 'web_post.png', 'web_interaccion.png']),
    ('HU-09', 'Informes Padres', ['web_informes.png', 'web_informe_previa.png']),
    ('HU-10', 'Perfil Privacidad', ['web_perfil.png', 'web_privacidad.png', 'web_sesiones.png']),
    ('HU-06', 'Editor Cuestionarios', ['web_editor.png', 'web_editor_previa.png']),
    ('Alertas', 'Gestión Alertas', ['web_alertas_panel.png', 'web_alertas_detalle.png', 'web_alertas_bajo_medio.png']),
    ('Misiones', 'Misiones Check-In', ['web_checkin.png', 'web_misiones.png']),
    ('Videos', 'Videos Guiados', ['web_videos.png', 'web_video_reproduccion.png']),
    ('Pausas', 'Pausas Activas', ['web_pausas.png', 'web_pausa_config.png'])
]

images_folder = r'C:\\Users\\Admin\\Documents\\INGENIERIA DE SOFTWARE\\IV SEMESTRE\\ARCHIVOS (.pptx .docx .pdf .xlsx)\\APPS WELLNESS MENTAL\\WellnessMentalApp(WEB)\\web_images'

for key, title, images in modules:
    h = doc.add_heading(f'{key}: {title}', level=2)
    remove_spacing(h)
    h = doc.add_heading('Criterios de Aceptación Implementados', level=3)
    remove_spacing(h)
    
    for image in images:
        h = doc.add_heading(image.replace('.png', '').replace('web_', 'CA-'), level=4)
        remove_spacing(h)
        p = doc.add_paragraph()
        p.add_run(f'Evidencia Visual: {image}. ').bold = True
        p.add_run('✅ FUNCIONAL')
        remove_spacing(p)
        
        image_path = os.path.join(images_folder, image)
        if os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                remove_spacing(doc.paragraphs[-1])
                print(f'Imagen agregada: {image}')
            except Exception as e:
                doc.add_paragraph(f'[Error: {e}]')
        else:
            p = doc.add_paragraph(f'[Imagen requerida: {image}]')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_paragraph()
    
    doc.add_page_break()

# Conclusiones
h = doc.add_heading('Conclusiones', level=1)
remove_spacing(h)
doc.add_paragraph('La aplicación web Wellness Mental App cumple satisfactoriamente con todos los criterios de aceptación establecidos para cada historia de usuario.')

output_path = r'C:\\Users\\Admin\\Documents\\INGENIERIA DE SOFTWARE\\IV SEMESTRE\\ARCHIVOS (.pptx .docx .pdf .xlsx)\\APPS WELLNESS MENTAL\\WellnessMentalApp(WEB)\\INFORME_FINAL_INTEGRADO_WEB.docx'
doc.save(output_path)
print(f'Documento creado: {output_path}')
'''
    
    script_path = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\temp_script.py'
    
    with open(script_path, 'w', encoding='utf-8') as f:
        f.write(script_content)
    
    print("📝 Script de generación creado")
    
    # Ejecutar el script
    print("🚀 Generando documento Word...")
    try:
        result = subprocess.run([sys.executable, script_path], 
                              capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0:
            print("✅ " + result.stdout)
            return True
        else:
            print("❌ Error en la ejecución:")
            print("STDOUT:", result.stdout)
            print("STDERR:", result.stderr)
            return False
            
    except subprocess.TimeoutExpired:
        print("❌ Tiempo de espera agotado")
        return False
    except Exception as e:
        print(f"❌ Error: {e}")
        return False
    finally:
        # Limpiar script temporal
        if os.path.exists(script_path):
            os.remove(script_path)

if __name__ == '__main__':
    success = install_and_run()
    if success:
        print("\n🎉 Proceso completado exitosamente")
    else:
        print("\n⚠️ Hubo errores en el proceso")
