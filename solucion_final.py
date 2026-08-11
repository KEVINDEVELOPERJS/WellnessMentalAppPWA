from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import os
import shutil

def create_final_solution():
    """Solución final con todas las imágenes y mejor calidad"""
    
    original_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp\INFORME_FINAL_INTEGRADO.docx'
    web_images_doc = r'C:\Users\Admin\Documents\FOTOS DE MODULOS APP WEB.docx'
    output_doc = r'C:\Users\Admin\Documents\INGENIERIA DE SOFTWARE\IV SEMESTRE\ARCHIVOS (.pptx .docx .pdf .xlsx)\APPS WELLNESS MENTAL\WellnessMentalApp(WEB)\INFORME_FINAL_INTEGRADO_WEB_SOLUCION_FINAL.docx'
    
    print("=== CREANDO SOLUCIÓN FINAL ===")
    print("1. Copiando documento original...")
    shutil.copy2(original_doc, output_doc)
    
    # Modificar texto
    replacements = {
        'Android': 'Web',
        'android': 'web',
        'APK': 'Web',
        'apk': 'web',
        'Android Studio': 'HTML5/CSS3/JavaScript',
        'Kotlin': 'JavaScript',
        'Java': 'JavaScript',
        'XML': 'HTML',
        'Gradle': 'NPM',
        'Room Database': 'IndexedDB',
        'SQLite': 'IndexedDB',
        'SharedPreferences': 'LocalStorage',
        'Retrofit': 'Fetch API',
        'Activity': 'Page',
        'aplicación móvil': 'aplicación web',
        'app móvil': 'app web',
        'móvil': 'web',
        'native': 'web',
        'nativo': 'web',
    }
    
    print("2. Modificando texto Android -> Web...")
    doc = Document(output_doc)
    
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)
    
    # Modificar título
    for paragraph in doc.paragraphs:
        if 'Android' in paragraph.text and ('INFORME' in paragraph.text or 'Funcionamiento' in paragraph.text):
            paragraph.text = paragraph.text.replace('Android', 'Web')
            if 'Wellness Mental App' in paragraph.text:
                paragraph.text = paragraph.text.replace('Wellness Mental App', 'Wellness Mental App Web')
    
    doc.save(output_doc)
    
    print("3. Procesando imágenes del documento web...")
    web_images = {}
    with zipfile.ZipFile(web_images_doc, 'r') as z:
        for f in z.namelist():
            if f.startswith('word/media/image'):
                num = f.replace('word/media/image', '').replace('.png', '')
                web_images[num] = z.read(f)
    
    print(f"   Total imágenes web: {len(web_images)}")
    
    print("4. Reemplazando imágenes en el documento...")
    temp = output_doc + '.tmp'
    
    with zipfile.ZipFile(output_doc, 'r') as z1:
        with zipfile.ZipFile(temp, 'w', zipfile.ZIP_DEFLATED) as z2:
            replaced = 0
            added = 0
            
            for info in z1.infolist():
                data = z1.read(info.filename)
                if info.filename.startswith('word/media/image'):
                    num = info.filename.replace('word/media/image', '').replace('.png', '')
                    if num in web_images:
                        z2.writestr(info.filename, web_images[num])
                        replaced += 1
                    else:
                        z2.writestr(info, data)
                else:
                    z2.writestr(info, data)
            
            # Añadir imágenes adicionales (35-44)
            for num in range(35, 45):
                num_str = str(num)
                if num_str in web_images:
                    new_filename = f'word/media/image{num_str}.png'
                    z2.writestr(new_filename, web_images[num_str])
                    added += 1
                    print(f"   Añadida imagen adicional: image{num_str}.png")
            
            print(f"   Imágenes reemplazadas: {replaced}")
            print(f"   Imágenes adicionales: {added}")
    
    os.remove(output_doc)
    shutil.move(temp, output_doc)
    
    print("5. Añadiendo imágenes adicionales al documento con python-docx...")
    doc = Document(output_doc)
    
    # Añadir página nueva para las imágenes adicionales
    doc.add_page_break()
    
    heading = doc.add_heading('Imágenes Adicionales de Módulos Web', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Añadir las imágenes adicionales (35-44) con buen tamaño
    for i in range(35, 45):
        num_str = str(i)
        if num_str in web_images:
            temp_img = f'temp_img_{num_str}.png'
            with open(temp_img, 'wb') as f:
                f.write(web_images[num_str])
            
            try:
                doc.add_heading(f'Imagen {num_str}', level=2)
                doc.add_picture(temp_img, width=Inches(6.5))  # Tamaño grande para buena visibilidad
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                print(f"   Imagen {num_str} añadida con tamaño 6.5\"")
            except Exception as e:
                print(f"   Error con imagen {num_str}: {e}")
            finally:
                if os.path.exists(temp_img):
                    os.remove(temp_img)
    
    doc.save(output_doc)
    
    print("6. Ajustando tamaño de imágenes existentes...")
    doc = Document(output_doc)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for shape in run._element.xpath('.//pic:pic'):
                for extent in shape.xpath('.//a:ext'):
                    # Ajustar a 6.5 pulgadas de ancho
                    extent.set('{http://schemas.openxmlformats.org/drawingml/2006/main}cx', str(int(Inches(6.5).emu)))
                    extent.set('{http://schemas.openxmlformats.org/drawingml/2006/main}cy', str(int(Inches(4.5).emu)))
    
    doc.save(output_doc)
    
    print(f"\n=== DOCUMENTO FINAL CREADO ===")
    print(f"Archivo: {output_doc}")
    print(f"- Texto modificado de Android a Web")
    print(f"- {replaced} imágenes reemplazadas (mismo tamaño que original)")
    print(f"- {added} imágenes adicionales añadidas (dashboard psicólogo)")
    print(f"- Todas las imágenes con tamaño ajustado a 6.5\" x 4.5\"")
    print(f"- Total imágenes: {replaced + added}")
    
    return output_doc

if __name__ == "__main__":
    create_final_solution()
